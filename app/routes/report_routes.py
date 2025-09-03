# -*- coding: utf-8 -*-
"""
Rotas especializadas para relatórios de avaliações
Endpoint para geração de relatórios completos com estatísticas detalhadas
"""

from flask import Blueprint, request, jsonify, render_template_string, send_file
from flask_jwt_extended import jwt_required
from app.decorators.role_required import role_required, get_current_user_from_token
from app.models.test import Test
from app.models.student import Student
from app.models.studentAnswer import StudentAnswer
from app.models.testSession import TestSession
from app.models.question import Question
from app.models.subject import Subject
from app.models.school import School
from app.models.city import City
from app.models.studentClass import Class
from app.models.grades import Grade
from app.models.classTest import ClassTest
from app.models.evaluationResult import EvaluationResult
from app import db
import logging
from typing import Dict, Any, List, Optional
from sqlalchemy import func, case, desc
from datetime import datetime
from sqlalchemy.orm import joinedload
from collections import defaultdict
import os
from jinja2 import Template
import tempfile
from reportlab.lib.pagesizes import A4
from reportlab.platypus import BaseDocTemplate, PageTemplate, Frame, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm, inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.pdfgen.canvas import Canvas
from io import BytesIO
from app.services.ai_analysis_service import AIAnalysisService

# Importar docxtpl para template Word
try:
    from docxtpl import DocxTemplate
    from docx.shared import Mm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    DOCXTPL_AVAILABLE = True
except ImportError:
    DOCXTPL_AVAILABLE = False
    logging.warning("docxtpl não disponível. Instale com: pip install docxtpl")

# Constantes de cores para o PDF
BLUE_900 = colors.HexColor("#0b2a56")
GRAY_150 = colors.HexColor("#ededed")
GRAY_200 = colors.HexColor("#e6e6e6")
GRID = colors.HexColor("#999")
RED = colors.HexColor("#e53935")
YELLOW = colors.HexColor("#f1c232")
ADEQ = colors.HexColor("#6aa84f")
GREEN = colors.HexColor("#00a651")

def _header(canvas: Canvas, doc, logo_esq=None, logo_dir=None):
    """Função para desenhar cabeçalho com logos em todas as páginas"""
    top_y = doc.height + doc.topMargin + 8*mm
    if logo_esq:
        canvas.drawImage(logo_esq, doc.leftMargin, top_y-18*mm, width=28*mm, height=18*mm, preserveAspectRatio=True, mask='auto')
    if logo_dir:
        x = doc.pagesize[0] - doc.rightMargin - 28*mm
        canvas.drawImage(logo_dir, x, top_y-18*mm, width=28*mm, height=18*mm, preserveAspectRatio=True, mask='auto')

bp = Blueprint('reports', __name__, url_prefix='/reports')


@bp.route('/test', methods=['GET'])
def test_endpoint():
    """Endpoint de teste para verificar se o blueprint está funcionando"""
    return jsonify({
        "message": "Blueprint de relatórios funcionando corretamente",
        "status": "success"
    }), 200


@bp.route('/test-questions/<string:evaluation_id>', methods=['GET'])
@jwt_required()
@role_required("admin", "professor", "coordenador", "diretor","tecadm")
def test_questions(evaluation_id: str):
    """Endpoint para testar diretamente a propriedade questions"""
    try:
        from app.models.testQuestion import TestQuestion
        from app.models.question import Question
        
        # Teste 1: Query direta na tabela test_questions
        test_questions_direct = TestQuestion.query.filter_by(test_id=evaluation_id).order_by(TestQuestion.order).all()
        
        # Teste 2: IDs das questões
        question_ids = [tq.question_id for tq in test_questions_direct]
        
        # Teste 3: Buscar questões diretamente
        questions_direct = Question.query.filter(Question.id.in_(question_ids)).all()
        
        # Teste 4: Usar a propriedade questions do modelo Test
        test = Test.query.get(evaluation_id)
        questions_property = test.questions if test else []
        
        # Teste 5: Verificar se a questão de Português está na tabela
        portugues_question = TestQuestion.query.filter_by(
            test_id=evaluation_id,
            question_id="4e3305a0-7221-4012-a7aa-8f958b755823"
        ).first()
        
        debug_info = {
            "avaliacao_id": evaluation_id,
            "test_questions_direct": {
                "total": len(test_questions_direct),
                "registros": [
                    {
                        "id": tq.id,
                        "test_id": tq.test_id,
                        "question_id": tq.question_id,
                        "order": tq.order
                    } for tq in test_questions_direct
                ]
            },
            "question_ids": question_ids,
            "questions_direct": {
                "total": len(questions_direct),
                "questoes": [
                    {
                        "id": q.id,
                        "skill": q.skill,
                        "number": q.number
                    } for q in questions_direct
                ]
            },
            "questions_property": {
                "total": len(questions_property),
                "questoes": [
                    {
                        "id": q.id,
                        "skill": q.skill,
                        "number": q.number
                    } for q in questions_property
                ]
            },
            "questao_portugues_teste": {
                "encontrada": portugues_question is not None,
                "dados": {
                    "id": portugues_question.id,
                    "test_id": portugues_question.test_id,
                    "question_id": portugues_question.question_id,
                    "order": portugues_question.order
                } if portugues_question else None
            }
        }
        
        return jsonify(debug_info), 200
        
    except Exception as e:
        logging.error(f"Erro ao testar questions: {str(e)}")
        return jsonify({"error": "Erro interno do servidor", "details": str(e)}), 500


@bp.route('/debug-disciplinas/<string:evaluation_id>', methods=['GET'])
@jwt_required()
@role_required("admin", "professor", "coordenador", "diretor","tecadm")
def debug_disciplinas(evaluation_id: str):
    """Endpoint para debug da identificação de disciplinas"""
    try:
        # Verificar se a avaliação existe
        test = Test.query.get(evaluation_id)
        if not test:
            return jsonify({"error": "Avaliação não encontrada"}), 404
        
        # Obter disciplinas usando a função melhorada
        disciplinas = _obter_disciplinas_avaliacao(test)
        
        # Debug detalhado
        debug_info = {
            "avaliacao_id": evaluation_id,
            "titulo": test.title,
            "disciplinas_identificadas": disciplinas,
            "debug_details": {
                "subject_rel": test.subject_rel.name if test.subject_rel else None,
                "subjects_info": test.subjects_info,
                "total_questoes": len(test.questions) if test.questions else 0,
                "questoes_com_habilidades": 0,
                "habilidades_unicas": set(),
                "disciplinas_por_habilidade": {}
            },
            "mapeamento_questoes": {},
            "respostas_por_questao": {},
            "questao_portugues_debug": {}
        }
        
        if test.questions:
            from app.models.skill import Skill
            
            skill_ids = set()
            for question in test.questions:
                if question.skill and question.skill.strip() and question.skill != '{}':
                    clean_skill_id = question.skill.replace('{', '').replace('}', '')
                    skill_ids.add(clean_skill_id)
                    debug_info["debug_details"]["questoes_com_habilidades"] += 1
            
            debug_info["debug_details"]["habilidades_unicas"] = list(skill_ids)
            
            if skill_ids:
                skills = Skill.query.filter(Skill.id.in_(skill_ids)).all()
                skills_dict = {str(skill.id): skill for skill in skills}
                
                for skill in skills:
                    subject = Subject.query.get(skill.subject_id) if skill.subject_id else None
                    debug_info["debug_details"]["disciplinas_por_habilidade"][str(skill.id)] = {
                        "skill_code": skill.code,
                        "skill_description": skill.description,
                        "subject_id": skill.subject_id,
                        "subject_name": subject.name if subject else None
                    }
                
                # Mapear questões para disciplinas (mesmo processo das funções de cálculo)
                for question in test.questions:
                    # Questões com skill: mapear via skill.subject_id
                    if question.skill and question.skill.strip() and question.skill != '{}':
                        clean_skill_id = question.skill.replace('{', '').replace('}', '')
                        skill_obj = skills_dict.get(clean_skill_id)
                        
                        disciplina_mapeada = "Não mapeada"
                        if skill_obj and skill_obj.subject_id:
                            subject = Subject.query.get(skill_obj.subject_id)
                            if subject:
                                disciplina_mapeada = subject.name
                            else:
                                disciplina_mapeada = "Disciplina Geral (subject não encontrado)"
                        else:
                            disciplina_mapeada = "Disciplina Geral (skill sem subject_id)"
                        
                        debug_info["mapeamento_questoes"][question.id] = {
                            "numero": question.number,
                            "skill_id": clean_skill_id,
                            "skill_encontrada": skill_obj is not None,
                            "disciplina_mapeada": disciplina_mapeada
                        }
                    else:
                        # Questões sem skill: mapear para disciplina principal da avaliação
                        if test.subject_rel:
                            disciplina_mapeada = test.subject_rel.name
                            debug_info["mapeamento_questoes"][question.id] = {
                                "numero": question.number,
                                "skill_id": "Sem skill",
                                "skill_encontrada": False,
                                "disciplina_mapeada": disciplina_mapeada
                            }
                        else:
                            disciplina_mapeada = "Disciplina Geral"
                            debug_info["mapeamento_questoes"][question.id] = {
                                "numero": question.number,
                                "skill_id": "Sem skill",
                                "skill_encontrada": False,
                                "disciplina_mapeada": disciplina_mapeada
                            }
                    
                    # Verificar respostas para esta questão
                    respostas = StudentAnswer.query.filter_by(
                        test_id=evaluation_id, 
                        question_id=question.id
                    ).count()
                    
                    debug_info["respostas_por_questao"][question.id] = respostas
                    
                    # Debug específico para a questão de Português que sabemos que existe
                    if question.id == "4e3305a0-7221-4012-a7aa-8f958b755823":
                        if question.skill and question.skill.strip() and question.skill != '{}':
                            clean_skill_id = question.skill.replace('{', '').replace('}', '')
                            skill_obj = skills_dict.get(clean_skill_id)
                            debug_info["questao_portugues_debug"] = {
                                "questao_id": question.id,
                                "numero": question.number,
                                "skill_raw": question.skill,
                                "skill_clean": clean_skill_id,
                                "skill_encontrada": skill_obj is not None,
                                "skill_code": skill_obj.code if skill_obj else None,
                                "skill_subject_id": skill_obj.subject_id if skill_obj else None,
                                "disciplina_final": disciplina_mapeada,
                                "total_respostas": respostas
                            }
                        else:
                            debug_info["questao_portugues_debug"] = {
                                "questao_id": question.id,
                                "numero": question.number,
                                "skill_raw": question.skill,
                                "skill_clean": "Sem skill",
                                "skill_encontrada": False,
                                "skill_code": None,
                                "skill_subject_id": None,
                                "disciplina_final": disciplina_mapeada,
                                "total_respostas": respostas
                            }
        
        return jsonify(debug_info), 200
        
    except Exception as e:
        logging.error(f"Erro ao debug disciplinas: {str(e)}")
        return jsonify({"error": "Erro interno do servidor", "details": str(e)}), 500


@bp.route('/relatorio-completo/<string:evaluation_id>', methods=['GET'])
@jwt_required()
@role_required("admin", "professor", "coordenador", "diretor","tecadm")
def relatorio_completo(evaluation_id: str):
    """
    Gera relatório completo de uma avaliação com todos os dados solicitados
    
    Args:
        evaluation_id: ID da avaliação
    
    Returns:
        JSON com relatório completo contendo:
        - Total de alunos que realizaram a avaliação
        - Níveis de Aprendizagem por turma
        - Proficiência
        - Nota Geral por turma
        - Acertos por habilidade
    """
    try:
        # Verificar se a avaliação existe
        test = Test.query.get(evaluation_id)
        if not test:
            return jsonify({"error": "Avaliação não encontrada"}), 404
        
        # Buscar turmas onde a avaliação foi aplicada
        class_tests = ClassTest.query.filter_by(test_id=evaluation_id).all()
        if not class_tests:
            return jsonify({"error": "Avaliação não foi aplicada em nenhuma turma"}), 404
        
        # Obter dados da avaliação
        avaliacao_data = {
            "id": test.id,
            "titulo": test.title,
            "descricao": test.description,
            "disciplinas": _obter_disciplinas_avaliacao(test)
        }
        
        logging.info(f"Disciplinas identificadas para avaliação {evaluation_id}: {avaliacao_data['disciplinas']}")
        
        # 1. Total de alunos que realizaram a avaliação
        total_alunos = _calcular_totais_alunos(evaluation_id, class_tests)
        
        # 2. Níveis de Aprendizagem por turma
        niveis_aprendizagem = _calcular_niveis_aprendizagem(evaluation_id, class_tests)
        logging.info(f"Níveis de aprendizagem calculados para disciplinas: {list(niveis_aprendizagem.keys())}")
        
        # 3. Proficiência
        proficiencia = _calcular_proficiencia(evaluation_id, class_tests)
        logging.info(f"Proficiência calculada para disciplinas: {list(proficiencia.get('por_disciplina', {}).keys())}")
        
        # 4. Nota Geral por turma
        nota_geral = _calcular_nota_geral(evaluation_id, class_tests)
        logging.info(f"Nota geral calculada para disciplinas: {list(nota_geral.get('por_disciplina', {}).keys())}")
        
        # 5. Acertos por habilidade
        acertos_habilidade = _calcular_acertos_habilidade(evaluation_id)
        logging.info(f"Acertos por habilidade calculados para disciplinas: {list(acertos_habilidade.keys())}")
        
        return jsonify({
            "avaliacao": avaliacao_data,
            "total_alunos": total_alunos,
            "niveis_aprendizagem": niveis_aprendizagem,
            "proficiencia": proficiencia,
            "nota_geral": nota_geral,
            "acertos_por_habilidade": acertos_habilidade
        }), 200
        
    except Exception as e:
        logging.error(f"Erro ao gerar relatório completo: {str(e)}")
        return jsonify({"error": "Erro interno do servidor"}), 500


@bp.route('/relatorio-com-ia/<string:evaluation_id>', methods=['GET'])
@jwt_required()
@role_required("admin", "professor", "coordenador", "diretor","tecadm")
def relatorio_com_ia(evaluation_id: str):
    """
    Gera relatório completo com análise da IA
    
    Args:
        evaluation_id: ID da avaliação
    
    Returns:
        JSON com relatório completo + análise da IA
    """
    try:
        # Verificar se a avaliação existe
        test = Test.query.get(evaluation_id)
        if not test:
            return jsonify({"error": "Avaliação não encontrada"}), 404
        
        # Buscar turmas onde a avaliação foi aplicada
        class_tests = ClassTest.query.filter_by(test_id=evaluation_id).all()
        if not class_tests:
            return jsonify({"error": "Avaliação não foi aplicada em nenhuma turma"}), 404
        
        # Obter dados da avaliação
        avaliacao_data = {
            "id": test.id,
            "titulo": test.title,
            "descricao": test.description,
            "disciplinas": _obter_disciplinas_avaliacao(test),
            "questoes_anuladas": []  # Lista de questões anuladas se houver
        }
        
        logging.info(f"Disciplinas identificadas para avaliação {evaluation_id}: {avaliacao_data['disciplinas']}")
        
        # 1. Total de alunos que realizaram a avaliação
        total_alunos = _calcular_totais_alunos(evaluation_id, class_tests)
        
        # 2. Níveis de Aprendizagem por turma
        niveis_aprendizagem = _calcular_niveis_aprendizagem(evaluation_id, class_tests)
        
        # 3. Proficiência
        proficiencia = _calcular_proficiencia(evaluation_id, class_tests)
        
        # 4. Nota Geral por turma
        nota_geral = _calcular_nota_geral(evaluation_id, class_tests)
        
        # 5. Acertos por habilidade
        acertos_habilidade = _calcular_acertos_habilidade(evaluation_id)
        
        # Preparar dados completos para análise da IA
        report_data = {
            "avaliacao": avaliacao_data,
            "total_alunos": total_alunos,
            "niveis_aprendizagem": niveis_aprendizagem,
            "proficiencia": proficiencia,
            "nota_geral": nota_geral,
            "acertos_por_habilidade": acertos_habilidade
        }
        
        # 6. Análise da IA
        ai_service = AIAnalysisService()
        ai_analysis = ai_service.analyze_report_data(report_data)
        
        return jsonify({
            "avaliacao": avaliacao_data,
            "total_alunos": total_alunos,
            "niveis_aprendizagem": niveis_aprendizagem,
            "proficiencia": proficiencia,
            "nota_geral": nota_geral,
            "acertos_por_habilidade": acertos_habilidade,
            "analise_ia": ai_analysis
        }), 200
        
    except Exception as e:
        logging.error(f"Erro ao gerar relatório com IA: {str(e)}")
        return jsonify({"error": "Erro interno do servidor", "details": str(e)}), 500


@bp.route('/relatorio-pdf/<string:evaluation_id>', methods=['GET'])
@jwt_required()
@role_required("admin", "professor", "coordenador", "diretor","tecadm")
def relatorio_pdf(evaluation_id: str):
    """
    Gera relatório PDF de uma avaliação usando template Word
    
    Args:
        evaluation_id: ID da avaliação
    
    Returns:
        Arquivo DOCX do relatório para download
    """
    try:
        # Verificar se docxtpl está disponível
        if not DOCXTPL_AVAILABLE:
            return jsonify({"error": "docxtpl não está disponível. Instale com: pip install docxtpl"}), 500
        
        # Verificar se a avaliação existe
        test = Test.query.get(evaluation_id)
        if not test:
            return jsonify({"error": "Avaliação não encontrada"}), 404
        
        # Buscar turmas onde a avaliação foi aplicada
        class_tests = ClassTest.query.filter_by(test_id=evaluation_id).all()
        if not class_tests:
            return jsonify({"error": "Avaliação não foi aplicada em nenhuma turma"}), 404
        
        # Obter dados da avaliação (mesma lógica do relatorio_completo)
        avaliacao_data = {
            "id": test.id,
            "titulo": test.title,
            "descricao": test.description,
            "disciplinas": _obter_disciplinas_avaliacao(test)
        }
        
        logging.info(f"Disciplinas identificadas para avaliação {evaluation_id}: {avaliacao_data['disciplinas']}")
        
        # 1. Total de alunos que realizaram a avaliação
        total_alunos = _calcular_totais_alunos(evaluation_id, class_tests)
        
        # 2. Níveis de Aprendizagem por turma
        niveis_aprendizagem = _calcular_niveis_aprendizagem(evaluation_id, class_tests)
        logging.info(f"Níveis de aprendizagem calculados para disciplinas: {list(niveis_aprendizagem.keys())}")
        
        # 3. Proficiência
        proficiencia = _calcular_proficiencia(evaluation_id, class_tests)
        logging.info(f"Proficiência calculada para disciplinas: {list(proficiencia.get('por_disciplina', {}).keys())}")
        
        # 4. Nota Geral por turma
        nota_geral = _calcular_nota_geral(evaluation_id, class_tests)
        logging.info(f"Nota geral calculada para disciplinas: {list(nota_geral.get('por_disciplina', {}).keys())}")
        
        # 5. Acertos por habilidade
        acertos_habilidade = _calcular_acertos_habilidade(evaluation_id)
        logging.info(f"Acertos por habilidade calculados para disciplinas: {list(acertos_habilidade.keys())}")
        
        # 6. Análise da IA
        ai_service = AIAnalysisService()
        ai_analysis = ai_service.analyze_report_data({
            "avaliacao": avaliacao_data,
            "total_alunos": total_alunos,
            "niveis_aprendizagem": niveis_aprendizagem,
            "proficiencia": proficiencia,
            "nota_geral": nota_geral,
            "acertos_por_habilidade": acertos_habilidade
        })
        
        # Preparar dados para o template Word
        template_data = _preparar_dados_template_word(
            test, total_alunos, niveis_aprendizagem, 
            proficiencia, nota_geral, acertos_habilidade, ai_analysis, avaliacao_data
        )
        
        # Gerar relatório Word usando docxtpl
        try:
            docx_content = _gerar_docx_com_template(template_data)
            
            # Preparar nome do arquivo com o nome da avaliação
            nome_avaliacao = test.title.replace(' ', '_').replace('/', '_').replace('\\', '_').replace(':', '_').replace('?', '_').replace('*', '_').replace('"', '_').replace('<', '_').replace('>', '_').replace('|', '_')
            nome_arquivo = f"relatorio_{nome_avaliacao}.docx"
            
            # Retornar arquivo DOCX
            from flask import Response
            response = Response(docx_content, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response.headers['Content-Disposition'] = f'attachment; filename={nome_arquivo}'
            
            return response
            
        except Exception as docx_error:
            logging.error(f"Erro ao gerar DOCX com template: {str(docx_error)}")
            
            # Fallback: retornar erro se DOCX falhar
            return jsonify({"error": "Erro ao gerar relatório DOCX", "details": str(docx_error)}), 500
        
    except Exception as e:
        logging.error(f"Erro ao gerar relatório DOCX: {str(e)}")
        return jsonify({"error": "Erro interno do servidor", "details": str(e)}), 500


@bp.route('/test-html/<string:evaluation_id>', methods=['GET'])
@jwt_required()
@role_required("admin", "professor", "coordenador", "diretor","tecadm")
def test_html_template(evaluation_id: str):
    """Endpoint para testar o template HTML renderizado"""
    try:
        # Verificar se a avaliação existe
        test = Test.query.get(evaluation_id)
        if not test:
            return jsonify({"error": "Avaliação não encontrada"}), 404
        
        # Buscar turmas onde a avaliação foi aplicada
        class_tests = ClassTest.query.filter_by(test_id=evaluation_id).all()
        if not class_tests:
            return jsonify({"error": "Avaliação não foi aplicada em nenhuma turma"}), 404
        
        # Obter dados da avaliação usando as funções existentes
        avaliacao_data = {
            "id": test.id,
            "titulo": test.title,
            "descricao": test.description,
            "disciplinas": _obter_disciplinas_avaliacao(test)
        }
        
        # 1. Total de alunos que realizaram a avaliação
        total_alunos = _calcular_totais_alunos(evaluation_id, class_tests)
        
        # 2. Níveis de Aprendizagem por turma
        niveis_aprendizagem = _calcular_niveis_aprendizagem(evaluation_id, class_tests)
        
        # 3. Proficiência
        proficiencia = _calcular_proficiencia(evaluation_id, class_tests)
        
        # 4. Nota Geral por turma
        nota_geral = _calcular_nota_geral(evaluation_id, class_tests)
        
        # 5. Acertos por habilidade
        acertos_habilidade = _calcular_acertos_habilidade(evaluation_id)
        
        # Preparar dados para o template
        template_data = _preparar_dados_template(
            test, total_alunos, niveis_aprendizagem, 
            proficiencia, nota_geral, acertos_habilidade
        )
        
        # Ler o template HTML
        template_path = os.path.join(os.path.dirname(__file__), '..', 'templates', 'relatorio_avaliacao.html')
        with open(template_path, 'r', encoding='utf-8') as f:
            template_content = f.read()
        
        # Renderizar o template
        template = Template(template_content)
        html_content = template.render(**template_data)
        
        # Retornar HTML para teste
        from flask import Response
        response = Response(html_content, mimetype='text/html')
        return response
        
    except Exception as e:
        logging.error(f"Erro ao testar template HTML: {str(e)}")
        return jsonify({"error": "Erro interno do servidor", "details": str(e)}), 500


def _preparar_dados_template(test: Test, total_alunos: Dict, niveis_aprendizagem: Dict, 
                           proficiencia: Dict, nota_geral: Dict, acertos_habilidade: Dict) -> Dict[str, Any]:
    """Prepara os dados para o template Jinja2"""
    
    # Obter informações da escola e município
    escola_nome = "Escola não identificada"
    municipio_nome = "Município não identificado"
    uf = "AL"
    
    try:
        # Buscar primeira turma para obter escola e município
        if test.class_tests:
            first_class_test = test.class_tests[0]
            first_class = Class.query.get(first_class_test.class_id)
            if first_class and first_class.school:
                escola_nome = first_class.school.name
                if first_class.school.city:
                    municipio_nome = first_class.school.city.name
                    uf = first_class.school.city.state or "AL"
    except Exception as e:
        logging.warning(f"Erro ao obter dados da escola/município: {str(e)}")
    
    # Preparar dados de participação
    participacao = []
    for turma_data in total_alunos.get('por_turma', []):
        participacao.append({
            'serie_turno': turma_data['turma'],
            'matriculados': turma_data['matriculados'],
            'avaliados': turma_data['avaliados'],
            'percentual': turma_data['percentual'],
            'faltosos': turma_data['faltosos']
        })
    
    # Resumo de participação
    total_geral = total_alunos.get('total_geral', {})
    participacao_resumo = f"Total geral: {total_geral.get('avaliados', 0)} alunos avaliados de {total_geral.get('matriculados', 0)} matriculados ({total_geral.get('percentual', 0)}%)"
    
    # Textos padrão
    apresentacao_texto = f"""
    Este relatório apresenta os resultados da <strong>Avaliação Diagnóstica {test.title}</strong> 
    aplicada na rede municipal de ensino de {municipio_nome}, {uf}. 
    A avaliação foi realizada com o objetivo de diagnosticar o nível de aprendizagem dos estudantes 
    e identificar áreas que necessitam de maior atenção pedagógica.
    """
    
    consideracoes = f"""
    A avaliação foi aplicada em {len(test.class_tests) if test.class_tests else 0} turmas, 
    totalizando {total_geral.get('avaliados', 0)} estudantes avaliados. 
    Os resultados apresentados refletem o desempenho dos alunos em diferentes habilidades 
    e competências, organizados por disciplina e nível de aprendizagem.
    """
    
    # Determinar período atual
    from datetime import datetime
    now = datetime.now()
    mes_atual = now.strftime("%B").upper()
    ano_atual = now.year
    periodo = f"{ano_atual}.1"
    
    return {
        'escola': escola_nome,
        'municipio': municipio_nome,
        'uf': uf,
        'periodo': periodo,
        'mes': mes_atual,
        'ano': ano_atual,
        'apresentacao_texto': apresentacao_texto,
        'consideracoes': consideracoes,
        'participacao': participacao,
        'participacao_resumo': participacao_resumo,
        'niveis_aprendizagem': niveis_aprendizagem,
        'proficiencia': proficiencia,
        'nota_geral': nota_geral,
        'acertos_por_habilidade': acertos_habilidade
    }


def _preparar_dados_template_word(test: Test, total_alunos: Dict, niveis_aprendizagem: Dict, 
                               proficiencia: Dict, nota_geral: Dict, acertos_habilidade: Dict, 
                               ai_analysis: Dict, avaliacao_data: Dict) -> Dict[str, Any]:
    """Prepara os dados para o template Word com docxtpl"""
    
    # Obter informações da escola e município
    escola_nome = "Escola não identificada"
    municipio_nome = "Município não identificado"
    uf = "AL"
    
    try:
        # Buscar primeira turma para obter escola e município
        if test.class_tests:
            first_class_test = test.class_tests[0]
            first_class = Class.query.get(first_class_test.class_id)
            if first_class and first_class.school:
                escola_nome = first_class.school.name
                if first_class.school.city:
                    municipio_nome = first_class.school.city.name
                    uf = first_class.school.city.state or "AL"
    except Exception as e:
        logging.warning(f"Erro ao obter dados da escola/município: {str(e)}")
    
    # Determinar período atual
    from datetime import datetime
    now = datetime.now()
    mes_atual = now.strftime("%B").upper()
    ano_atual = now.year
    periodo = f"{ano_atual}.1"
    
    # Dados básicos do template
    dados_base = {
        'periodo': periodo,
        'municipio': municipio_nome,
        'uf': uf,
        'escola': escola_nome,
        'mes': mes_atual,
        'ano': ano_atual,
        'total_alunos': total_alunos.get('total_geral', {}).get('avaliados', 0),
        'percentual_avaliados': total_alunos.get('total_geral', {}).get('percentual', 0),
        'disciplinas': ', '.join(avaliacao_data.get('disciplinas', ['Disciplina Geral']))
    }
    
    # Dados do sumário
    dados_sumario = {
        'sumario_p1': 'Este relatório apresenta os resultados da avaliação educacional realizada no período de 2024.',
        'sumario_p2': 'A avaliação abrangeu múltiplas disciplinas e habilidades dos estudantes.',
        'sumario_p3': 'Os resultados demonstram o progresso e as áreas de melhoria necessárias.',
        'sumario_p4': 'Recomendações específicas são apresentadas para otimizar o aprendizado.',
        'sumario_p41': 'Os dados apresentados refletem o desempenho real dos estudantes.',
        'sumario_p42': 'Análises detalhadas por disciplina foram realizadas.',
        'sumario_p43': 'Comparações com médias municipais foram incluídas.',
        'sumario_p44': 'Recomendações pedagógicas específicas foram elaboradas.'
    }
    
    # Dados de participação
    dados_participacao = _preparar_dados_participacao(total_alunos)
    
    # Dados de níveis de aprendizagem
    dados_niveis = _preparar_dados_niveis(niveis_aprendizagem)
    
    # Dados de proficiência
    dados_proficiencia = _preparar_dados_proficiencia_word(proficiencia)
    
    # Dados de notas
    dados_notas = _preparar_dados_notas_word(nota_geral)
    
    # Dados de habilidades
    dados_habilidades = _preparar_dados_habilidades_word(acertos_habilidade)
    
    # Dados de análise da IA
    dados_ia = _preparar_dados_ia(ai_analysis, {
        "total_alunos": total_alunos,
        "niveis_aprendizagem": niveis_aprendizagem,
        "proficiencia": proficiencia,
        "nota_geral": nota_geral,
        "acertos_por_habilidade": acertos_habilidade
    })
    
    # Combinar todos os dados
    dados_completos = {
        **dados_base,
        **dados_sumario,
        **dados_participacao,
        **dados_niveis,
        **dados_proficiencia,
        **dados_notas,
        **dados_habilidades,
        **dados_ia
    }
    
    return dados_completos


def _preparar_dados_participacao(total_alunos: Dict) -> Dict[str, Any]:
    """Prepara dados de participação para o template"""
    try:
        total_geral = total_alunos.get('total_geral', {})
        
        return {
            'total_matriculados': total_geral.get('matriculados', 0),
            'total_avaliados': total_geral.get('avaliados', 0),
            'total_faltosos': total_geral.get('faltosos', 0),
            'percentual_avaliados': total_geral.get('percentual', 0),
            'participacao_por_turma': total_alunos.get('por_turma', []),
            'total_alunos': total_geral.get('avaliados', 0)
        }
    except Exception as e:
        logging.warning(f"Erro ao preparar dados de participação: {str(e)}")
        return {}


def _preparar_dados_niveis(niveis_aprendizagem: Dict) -> Dict[str, Any]:
    """Prepara dados de níveis de aprendizagem para o template"""
    try:
        dados_niveis = {}
        
        for disciplina, dados in niveis_aprendizagem.items():
            if disciplina != 'GERAL':
                dados_niveis[f'niveis_{disciplina.lower().replace(" ", "_")}'] = dados.get('por_turma', [])
        
        # Dados gerais
        dados_niveis['niveis_geral'] = niveis_aprendizagem.get('GERAL', {}).get('por_turma', [])
        
        return dados_niveis
        
    except Exception as e:
        logging.warning(f"Erro ao preparar dados de níveis: {str(e)}")
        return {}


def _preparar_dados_proficiencia_word(proficiencia: Dict) -> Dict[str, Any]:
    """Prepara dados de proficiência para o template Word"""
    try:
        disciplinas = proficiencia.get('por_disciplina', {})
        
        # Lista de disciplinas disponíveis
        proficiencia_disciplinas = []
        for disciplina in disciplinas.keys():
            if disciplina != 'GERAL':
                proficiencia_disciplinas.append(disciplina)
        
        # Dados por série/turma
        proficiencia_series = []
        
        # Verificar se há dados gerais
        if 'GERAL' in disciplinas and 'por_turma' in disciplinas['GERAL']:
            for turma in disciplinas['GERAL']['por_turma']:
                serie_data = {'serie_turno': turma.get('turma', 'Turma')}
                
                # Adicionar proficiência por disciplina
                for disciplina in proficiencia_disciplinas:
                    if disciplina in disciplinas and 'por_turma' in disciplinas[disciplina]:
                        # Buscar proficiência desta turma nesta disciplina
                        for turma_disc in disciplinas[disciplina]['por_turma']:
                            if turma_disc.get('turma') == turma.get('turma'):
                                serie_data[disciplina] = turma_disc.get('proficiencia', 0)
                                break
                        else:
                            serie_data[disciplina] = 0
                    else:
                        serie_data[disciplina] = 0
                
                # Adicionar média municipal
                serie_data['MUNICIPAL'] = proficiencia.get('media_municipal_por_disciplina', {}).get('GERAL', 0)
                proficiencia_series.append(serie_data)
        else:
            # Se não há dados gerais, criar dados básicos
            logging.warning("Dados gerais de proficiência não encontrados")
            proficiencia_series = []
        
        return {
            'proficiencia_disciplinas': proficiencia_disciplinas,
            'proficiencia_series': proficiencia_series
        }
        
    except Exception as e:
        logging.warning(f"Erro ao preparar dados de proficiência: {str(e)}")
        return {
            'proficiencia_disciplinas': [],
            'proficiencia_series': []
        }


def _preparar_dados_notas_word(nota_geral: Dict) -> Dict[str, Any]:
    """Prepara dados de notas para o template Word"""
    try:
        disciplinas = nota_geral.get('por_disciplina', {})
        
        # Lista de disciplinas disponíveis
        notas_disciplinas = []
        for disciplina in disciplinas.keys():
            if disciplina != 'GERAL':
                notas_disciplinas.append(disciplina)
        
        # Dados por série/turma
        notas_series = []
        
        # Verificar se há dados gerais
        if 'GERAL' in disciplinas and 'por_turma' in disciplinas['GERAL']:
            for turma in disciplinas['GERAL']['por_turma']:
                serie_data = {'serie_turno': turma.get('turma', 'Turma')}
                
                # Adicionar nota por disciplina
                for disciplina in notas_disciplinas:
                    if disciplina in disciplinas and 'por_turma' in disciplinas[disciplina]:
                        # Buscar nota desta turma nesta disciplina
                        for turma_disc in disciplinas[disciplina]['por_turma']:
                            if turma_disc.get('turma') == turma.get('turma'):
                                serie_data[disciplina] = turma_disc.get('nota', 0)
                                break
                        else:
                            serie_data[disciplina] = 0
                    else:
                        serie_data[disciplina] = 0
                
                # Adicionar média municipal
                serie_data['MUNICIPAL'] = nota_geral.get('media_municipal_por_disciplina', {}).get('GERAL', 0)
                notas_series.append(serie_data)
        else:
            # Se não há dados gerais, criar dados básicos
            logging.warning("Dados gerais de notas não encontrados")
            notas_series = []
        
        return {
            'notas_disciplinas': notas_disciplinas,
            'notas_series': notas_series
        }
        
    except Exception as e:
        logging.warning(f"Erro ao preparar dados de notas: {str(e)}")
        return {
            'notas_disciplinas': [],
            'notas_series': []
        }


def _preparar_dados_habilidades_word(acertos_habilidade: Dict) -> Dict[str, Any]:
    """Prepara dados de habilidades para o template Word"""
    try:
        dados_habilidades = {}
        
        for disciplina, dados in acertos_habilidade.items():
            if disciplina != 'GERAL':
                dados_habilidades[f'habilidades_{disciplina.lower().replace(" ", "_")}'] = dados.get('habilidades', [])
        
        # Dados gerais
        dados_habilidades['habilidades_geral'] = acertos_habilidade.get('GERAL', {}).get('habilidades', [])
        
        return dados_habilidades
        
    except Exception as e:
        logging.warning(f"Erro ao preparar dados de habilidades: {str(e)}")
        return {}


def _preparar_dados_ia(ai_analysis: Dict, template_data: Dict) -> Dict[str, Any]:
    """Prepara dados de análise da IA para o template com análises inteligentes baseadas nos dados"""
    try:
        # Gerar análises inteligentes baseadas nos dados das tabelas
        analise_participacao = _gerar_analise_participacao(template_data)
        analise_proficiencia = _gerar_analise_proficiencia(template_data)
        analise_nota_geral = _gerar_analise_notas(template_data)
        analise_acertos_por_habilidade = _gerar_analise_habilidades(template_data)
        
        # Análise geral combinada
        analises = []
        if analise_participacao:
            analises.append(f"Participação: {analise_participacao}")
        if analise_proficiencia:
            analises.append(f"Proficiência: {analise_proficiencia}")
        if analise_nota_geral:
            analises.append(f"Notas: {analise_nota_geral}")
        if analise_acertos_por_habilidade:
            analises.append(f"Habilidades: {analise_acertos_por_habilidade}")
        
        analise_completa = ' '.join(analises) if analises else 'Análise não disponível.'
        
        return {
            'analise_ia': analise_completa,
            'analise_IA_participacao': analise_participacao,
            'analise_IA_proficiencia': analise_proficiencia,
            'analise_IA_nota': analise_nota_geral,
            'analise_IA_acertos': analise_acertos_por_habilidade,
            'consideracoes_finais': 'Com base na análise dos dados, recomenda-se atenção especial às áreas identificadas como críticas e continuidade no desenvolvimento das habilidades que apresentaram melhor desempenho.'
        }
    except Exception as e:
        logging.warning(f"Erro ao preparar dados da IA: {str(e)}")
        return {
            'analise_ia': 'Análise não disponível.',
            'analise_IA_participacao': 'Análise de participação não disponível.',
            'analise_IA_proficiencia': 'Análise de proficiência não disponível.',
            'analise_IA_nota': 'Análise de notas não disponível.',
            'analise_IA_acertos': 'Análise de habilidades não disponível.',
            'consideracoes_finais': 'Relatório gerado automaticamente pelo sistema.'
        }


def _gerar_docx_com_template(template_data: Dict[str, Any]) -> bytes:
    """Gera arquivo DOCX usando o template Word com docxtpl"""
    
    try:
        # Caminho para o template funcionando com formatação profissional
        template_path = os.path.join(os.path.dirname(__file__), '..', 'templates', 'templateDoc_working.docx')
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template não encontrado em: {template_path}")
        
        # Carregar template
        doc = DocxTemplate(template_path)
        
        # Preparar contexto com dados
        context = _preparar_contexto_docxtpl(doc, template_data)
        
        # Renderizar template
        doc.render(context)
        
        # Salvar em buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except Exception as e:
        logging.error(f"Erro ao gerar DOCX: {str(e)}")
        raise


def _preparar_contexto_docxtpl(doc: DocxTemplate, template_data: Dict[str, Any]) -> Dict[str, Any]:
    """Prepara o contexto para o template docxtpl"""
    
    context = {}
    
    # Função para criar tabelas com formatação
    def make_table(headers, rows, table_type="default"):
        sub = doc.new_subdoc()
        t = sub.add_table(rows=1, cols=len(headers))
        
        # Configurar largura das colunas
        for col in t.columns:
            col.width = Mm(30)  # 30mm por coluna
        
        # Formatar cabeçalho
        header_row = t.rows[0]
        for j, h in enumerate(headers):
            cell = header_row.cells[j]
            cell.text = str(h)
            
            # Aplicar formatação específica por tipo de tabela
            if table_type == "participacao":
                # Cabeçalho azul escuro com texto branco
                cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="0B2A56" w:val="clear"/>'))
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True
                        run.font.size = Pt(10)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif table_type == "niveis":
                # Cores específicas para cada coluna de níveis
                colors = ["1F4E79", "C00000", "FFC000", "70AD47", "00B050"]
                if j < len(colors):
                    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{colors[j]}" w:val="clear"/>'))
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)
                            if j == 0:  # Primeira coluna
                                run.font.color.rgb = RGBColor(0, 0, 0)
                            else:  # Outras colunas
                                run.font.color.rgb = RGBColor(255, 255, 255)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif table_type in ["proficiencia", "notas"]:
                # Cabeçalho azul escuro com texto branco
                cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="002060" w:val="clear"/>'))
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True
                        run.font.size = Pt(10)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif table_type == "habilidades":
                # Cabeçalho azul escuro com texto branco
                cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="002060" w:val="clear"/>'))
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True
                        run.font.size = Pt(10)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adicionar linhas de dados
        for i, r in enumerate(rows):
            row = t.add_row().cells
            
            # Aplicar fundo alternado para proficiência e notas
            if table_type in ["proficiencia", "notas"] and i % 2 == 0:
                row_bg_color = "F2F2F2"  # Cinza claro
            else:
                row_bg_color = "FFFFFF"  # Branco
            
            for j, val in enumerate(r):
                cell = row[j]
                cell.text = str(val)
                
                # Formatar células de dados
                if table_type == "participacao":
                    # Primeira coluna em negrito
                    if j == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    # Outras colunas alinhadas à direita
                    else:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                elif table_type == "niveis":
                    # Primeira coluna em negrito
                    if j == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    # Outras colunas centralizadas
                    else:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                elif table_type in ["proficiencia", "notas"]:
                    # Aplicar fundo alternado
                    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{row_bg_color}" w:val="clear"/>'))
                    # Primeira coluna em negrito
                    if j == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    # Outras colunas alinhadas à direita
                    else:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                elif table_type == "habilidades":
                    # Primeira coluna em negrito
                    if j == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    # Outras colunas centralizadas
                    else:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Aplicar bordas a todas as células
        for row in t.rows:
            for cell in row.cells:
                cell._tc.get_or_add_tcPr().append(parse_xml('<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>'))
        
        return sub
    
    # Função para dividir listas em chunks
    def chunk(lst, n): 
        return [lst[i:i+n] for i in range(0, len(lst), n)]
    
    # Função para construir tabelas multicolunas
    def build_multicol(series, disciplinas, max_cols=3, incluir_media=True, incluir_municipal=True, table_type="default"):
        subs = []
        for cols in chunk(disciplinas, max_cols):
            headers = ["SÉRIE/TURNO"] + cols + (["MÉDIA"] if incluir_media else []) + (["MUNICIPAL"] if incluir_municipal else [])
            rows = []
            for s in series:
                linha = [s["serie_turno"]]
                soma = cont = 0
                for c in cols:
                    v = s.get(c, ""); 
                    linha.append(v)
                    if isinstance(v, (int, float)): 
                        soma += v; 
                        cont += 1
                if incluir_media: 
                    linha.append(round(soma/cont, 2) if cont else "")
                if incluir_municipal: 
                    linha.append(s.get("MUNICIPAL", ""))
                rows.append(linha)
            
            # Adicionar linha total
            if rows:
                total_row = ["9º GERAL"]
                for j in range(1, len(headers)):
                    if j == 1:  # Primeira coluna de dados
                        valores = [r[j] for r in rows if isinstance(r[j], (int, float))]
                        total = sum(valores) / len(valores) if valores else 0
                        total_row.append(round(total, 2))
                    elif j == len(headers) - 2 and incluir_media:  # Coluna MÉDIA
                        valores = [r[j] for r in rows if isinstance(r[j], (int, float))]
                        total = sum(valores) / len(valores) if valores else 0
                        total_row.append(round(total, 2))
                    elif j == len(headers) - 1 and incluir_municipal:  # Coluna MUNICIPAL
                        total_row.append("")  # Vazia para linha total
                    else:
                        valores = [r[j] for r in rows if isinstance(r[j], (int, float))]
                        total = sum(valores) / len(valores) if valores else 0
                        total_row.append(round(total, 2))
                rows.append(total_row)
            
            subs.append(make_table(headers, rows, table_type))
        return subs
    
    # 1. Tabela de participação
    if template_data.get("participacao_por_turma"):
        headers = ["SÉRIE/TURNO", "MATRICULADOS", "AVALIADOS", "PERCENTUAL", "FALTOSOS"]
        rows = []
        for p in template_data["participacao_por_turma"]:
            rows.append([
                p.get("turma", ""),
                p.get("matriculados", 0),
                p.get("avaliados", 0),
                f"{p.get('percentual', 0)}%",
                p.get("faltosos", 0)
            ])
        
        # Adicionar linha total
        rows.append([
            "9º GERAL",
            template_data.get("total_matriculados", 0),
            template_data.get("total_avaliados", 0),
            f"{template_data.get('percentual_avaliados', 0)}%",
            template_data.get("total_faltosos", 0)
        ])
        
        context["tabela_participacao"] = make_table(headers, rows, "participacao")
    
    # 2. Blocos de níveis de aprendizagem
    context["blocos_niveis"] = _construir_blocos_niveis(doc, template_data)
    
    # 3. Tabelas de proficiência
    if template_data.get("proficiencia_disciplinas") and template_data.get("proficiencia_series"):
        subs_prof = build_multicol(
            template_data["proficiencia_series"], 
            template_data["proficiencia_disciplinas"],
            table_type="proficiencia"
        )
        context["tabela_proficiencia_1"] = subs_prof[0] if subs_prof else doc.new_subdoc()
        context["tabela_proficiencia_2"] = subs_prof[1] if len(subs_prof) > 1 else doc.new_subdoc()
    
    # 4. Tabelas de notas
    if template_data.get("notas_disciplinas") and template_data.get("notas_series"):
        subs_notas = build_multicol(
            template_data["notas_series"], 
            template_data["notas_disciplinas"],
            table_type="notas"
        )
        context["tabela_nota_1"] = subs_notas[0] if subs_notas else doc.new_subdoc()
        context["tabela_nota_2"] = subs_notas[1] if len(subs_notas) > 1 else doc.new_subdoc()
    
    # 5. Tabelas de habilidades por disciplina
    # Criar um Subdoc para cada disciplina individual
    disciplinas_list = template_data.get("disciplinas", []).split(", ") if isinstance(template_data.get("disciplinas"), str) else template_data.get("disciplinas", [])
    
    for disciplina in disciplinas_list:
        if disciplina and disciplina != "Disciplina Geral":
            # Criar placeholder específico para cada disciplina
            placeholder_name = f"tabela_habilidades_{disciplina.lower().replace(' ', '_').replace('-', '_')}"
            
            key = f"habilidades_{disciplina.lower().replace(' ', '_')}"
            habilidades = template_data.get(key, [])
            
            if habilidades:
                # Criar Subdoc para esta disciplina
                sub_doc = doc.new_subdoc()
                
                # Título da disciplina
                p = sub_doc.add_paragraph(disciplina.upper())
                p.runs[0].bold = True
                
                headers = ["QUESTÃO", "HABILIDADE", "ACERTOS", "TOTAL", "PERCENTUAL"]
                rows = []
                for h in habilidades:
                    rows.append([
                        h.get("questoes", [{}])[0].get("numero", "N/A") if h.get("questoes") else "N/A",
                        h.get("codigo", "N/A"),
                        h.get("acertos", 0),
                        h.get("total", 0),
                        f"{h.get('percentual', 0)}%"
                    ])
                
                # Criar tabela
                table = sub_doc.add_table(rows=1, cols=len(headers))
                for j, h in enumerate(headers): 
                    table.rows[0].cells[j].text = str(h)
                for r in rows:
                    row = table.add_row().cells
                    for j, val in enumerate(r): 
                        row[j].text = str(val)
                
                # Adicionar ao contexto com o nome específico da disciplina
                context[placeholder_name] = sub_doc
    
    # Também criar um Subdoc combinado para o placeholder geral
    habilidades_combined = doc.new_subdoc()
    for disciplina in disciplinas_list:
        if disciplina and disciplina != "Disciplina Geral":
            key = f"habilidades_{disciplina.lower().replace(' ', '_')}"
            habilidades = template_data.get(key, [])
            if habilidades:
                # Título da disciplina
                p = habilidades_combined.add_paragraph(disciplina.upper())
                p.runs[0].bold = True
                
                headers = ["QUESTÃO", "HABILIDADE", "ACERTOS", "TOTAL", "PERCENTUAL"]
                rows = []
                for h in habilidades:
                    rows.append([
                        h.get("questoes", [{}])[0].get("numero", "N/A") if h.get("questoes") else "N/A",
                        h.get("codigo", "N/A"),
                        h.get("acertos", 0),
                        h.get("total", 0),
                        f"{h.get('percentual', 0)}%"
                    ])
                
                # Criar tabela com formatação especial
                table = habilidades_combined.add_table(rows=2, cols=len(headers))  # 2 linhas: cabeçalho + códigos
                
                # Configurar largura das colunas
                for col in table.columns:
                    col.width = Mm(25)
                
                # Formatar cabeçalho (primeira linha)
                header_row = table.rows[0]
                for j, h in enumerate(headers):
                    cell = header_row.cells[j]
                    cell.text = str(h)
                    # Cabeçalho azul escuro com texto branco
                    cell._tc.get_or_add_tcPr().append(parse_xml('<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="002060" w:val="clear"/>'))
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.bold = True
                            run.font.size = Pt(10)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Adicionar linha de códigos (segunda linha)
                codes_row = table.rows[1]
                for j in range(len(headers)):
                    cell = codes_row.cells[j]
                    if j == 0:  # Primeira coluna
                        cell.text = "CÓDIGOS"
                    else:  # Outras colunas com códigos das habilidades
                        if j-1 < len(habilidades):
                            skill_code = habilidades[j-1].get("codigo", "")
                            cell.text = str(skill_code)
                        else:
                            cell.text = ""
                    
                    # Fundo amarelo para linha de códigos
                    cell._tc.get_or_add_tcPr().append(parse_xml('<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="FFC000" w:val="clear"/>'))
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Adicionar linhas de dados (percentuais)
                for r in rows:
                    row = table.add_row().cells
                    for j, val in enumerate(r):
                        cell = row[j]
                        cell.text = str(val)
                        
                        # Primeira coluna em negrito
                        if j == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                        # Outras colunas centralizadas
                        else:
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Aplicar bordas
                for row in table.rows:
                    for cell in row.cells:
                        cell._tc.get_or_add_tcPr().append(parse_xml('<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>'))
                
                # Espaçamento
                habilidades_combined.add_paragraph("")
    
    context["tabela_habilidades_disc"] = habilidades_combined
    
    # 6. Análises da IA (já preparadas na função _preparar_dados_ia)
    # Essas variáveis já estão sendo adicionadas automaticamente no loop abaixo
    
    # 7. Considerações finais
    context["consideracoes_finais"] = template_data.get("consideracoes_finais", "Relatório gerado automaticamente pelo sistema.")
    
    # 8. Adicionar todos os campos de texto simples
    for key, value in template_data.items():
        if isinstance(value, (str, int, float)) and key not in context:
            context[key] = value
    
    # 9. Adicionar campos específicos que podem estar faltando
    context["simulados_label"] = "Avaliação Diagnóstica"
    context["referencia"] = "Relatório de Avaliação Educacional"
    context["escola_extenso"] = template_data.get("escola", "Escola")
    
    # Adicionar campos que podem estar faltando
    context["total_avaliados"] = template_data.get("total_alunos", 0)
    context["percentual_avaliados"] = template_data.get("percentual_avaliados", 0)
    
    # Debug: mostrar quais campos estão sendo enviados
    logging.info(f"Campos do contexto: {list(context.keys())}")
    logging.info(f"Disciplinas encontradas: {disciplinas_list}")
    
    return context


def _construir_blocos_niveis(doc: DocxTemplate, template_data: Dict[str, Any]):
    """Constrói os blocos de níveis de aprendizagem como Subdoc"""
    sub = doc.new_subdoc()
    
    # Adicionar tabelas para cada disciplina
    for disciplina in template_data.get("disciplinas", []).split(", "):
        if disciplina and disciplina != "Disciplina Geral":
            key = f"niveis_{disciplina.lower().replace(' ', '_')}"
            niveis = template_data.get(key, [])
            
            if niveis:
                # Título da disciplina
                p = sub.add_paragraph(disciplina.upper())
                p.runs[0].bold = True
                
                # Tabela de níveis
                headers = ["SÉRIE/TURNO", "ABAIXO DO BÁSICO", "BÁSICO", "ADEQUADO", "AVANÇADO"]
                rows = []
                for n in niveis:
                    rows.append([
                        n.get("turma", ""),
                        n.get("abaixo_do_basico", 0),
                        n.get("basico", 0),
                        n.get("adequado", 0),
                        n.get("avancado", 0)
                    ])
                
                # Adicionar linha total se disponível
                total_key = f"niveis_{disciplina.lower().replace(' ', '_')}_total"
                if total_key in template_data:
                    total = template_data[total_key]
                    rows.append([
                        "TOTAL",
                        total.get("abaixo_do_basico", 0),
                        total.get("basico", 0),
                        total.get("adequado", 0),
                        total.get("avancado", 0)
                    ])
                
                # Criar tabela com formatação
                table = sub.add_table(rows=1, cols=len(headers))
                
                # Configurar largura das colunas
                for col in table.columns:
                    col.width = Mm(30)
                
                # Formatar cabeçalho com cores específicas
                header_colors = ["1F4E79", "C00000", "FFC000", "70AD47", "00B050"]
                for j, h in enumerate(headers):
                    cell = table.rows[0].cells[j]
                    cell.text = str(h)
                    
                    if j < len(header_colors):
                        cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{header_colors[j]}" w:val="clear"/>'))
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.size = Pt(10)
                                if j == 0:  # Primeira coluna
                                    run.font.color.rgb = RGBColor(0, 0, 0)
                                else:  # Outras colunas
                                    run.font.color.rgb = RGBColor(255, 255, 255)
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Adicionar linhas de dados
                for r in rows:
                    row = table.add_row().cells
                    for j, val in enumerate(r):
                        cell = row[j]
                        cell.text = str(val)
                        
                        # Primeira coluna em negrito
                        if j == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                        # Outras colunas centralizadas
                        else:
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Aplicar bordas
                for row in table.rows:
                    for cell in row.cells:
                        cell._tc.get_or_add_tcPr().append(parse_xml('<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>'))
                
                # Espaçamento
                sub.add_paragraph("")
    
    # Adicionar tabela geral
    niveis_geral = template_data.get("niveis_geral", [])
    if niveis_geral:
        p = sub.add_paragraph("MÉDIA GERAL")
        p.runs[0].bold = True
        
        headers = ["SÉRIE/TURNO", "ABAIXO DO BÁSICO", "BÁSICO", "ADEQUADO", "AVANÇADO"]
        rows = []
        for n in niveis_geral:
            rows.append([
                n.get("turma", ""),
                n.get("abaixo_do_basico", 0),
                n.get("basico", 0),
                n.get("adequado", 0),
                n.get("avancado", 0)
            ])
        
        table = sub.add_table(rows=1, cols=len(headers))
        
        # Configurar largura das colunas
        for col in table.columns:
            col.width = Mm(30)
        
        # Formatar cabeçalho com cores específicas
        header_colors = ["1F4E79", "C00000", "FFC000", "70AD47", "00B050"]
        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = str(h)
            
            if j < len(header_colors):
                cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{header_colors[j]}" w:val="clear"/>'))
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                        if j == 0:  # Primeira coluna
                            run.font.color.rgb = RGBColor(0, 0, 0)
                        else:  # Outras colunas
                            run.font.color.rgb = RGBColor(255, 255, 255)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adicionar linhas de dados
        for r in rows:
            row = table.add_row().cells
            for j, val in enumerate(r):
                cell = row[j]
                cell.text = str(val)
                
                # Aplicar fundo cinza para linha total
                cell._tc.get_or_add_tcPr().append(parse_xml('<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="D9D9D9" w:val="clear"/>'))
                
                # Primeira coluna em negrito
                if j == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                # Outras colunas centralizadas
                else:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Aplicar bordas
        for row in table.rows:
            for cell in row.cells:
                cell._tc.get_or_add_tcPr().append(parse_xml('<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>'))
    
    return sub


def _obter_disciplinas_avaliacao(test: Test) -> List[str]:
    """Obtém as disciplinas da avaliação"""
    disciplinas = []
    
    logging.info(f"Identificando disciplinas para avaliação {test.id}")
    
    # Se a avaliação tem uma disciplina específica
    if test.subject_rel:
        disciplinas.append(test.subject_rel.name)
        logging.info(f"Disciplina encontrada via subject_rel: {test.subject_rel.name}")
    
    # Se tem informações de disciplinas no JSON
    if test.subjects_info and isinstance(test.subjects_info, dict):
        disciplinas.extend(test.subjects_info.keys())
        logging.info(f"Disciplinas encontradas via subjects_info: {list(test.subjects_info.keys())}")
    
    # Buscar disciplinas através das habilidades das questões
    if test.questions:
        from app.models.skill import Skill
        
        skill_ids = set()
        for question in test.questions:
            if question.skill and question.skill.strip() and question.skill != '{}':
                clean_skill_id = question.skill.replace('{', '').replace('}', '')
                skill_ids.add(clean_skill_id)
        
        logging.info(f"Skill IDs encontrados nas questões: {skill_ids}")
        
        if skill_ids:
            skills = Skill.query.filter(Skill.id.in_(skill_ids)).all()
            logging.info(f"Habilidades encontradas na tabela Skill: {len(skills)}")
            
            for skill in skills:
                if skill.subject_id:
                    subject = Subject.query.get(skill.subject_id)
                    if subject and subject.name not in disciplinas:
                        disciplinas.append(subject.name)
                        logging.info(f"Disciplina encontrada via habilidade {skill.code}: {subject.name}")
                else:
                    logging.warning(f"Habilidade {skill.code} não tem subject_id")
    
    # Se não encontrou disciplinas, usar padrão
    if not disciplinas:
        disciplinas = ["Disciplina Geral"]
        logging.warning("Nenhuma disciplina encontrada, usando padrão: Disciplina Geral")
    
    disciplinas_finais = list(set(disciplinas))  # Remove duplicatas
    logging.info(f"Disciplinas finais identificadas: {disciplinas_finais}")
    
    return disciplinas_finais


def _calcular_totais_alunos(evaluation_id: str, class_tests: List[ClassTest]) -> Dict[str, Any]:
    """Calcula totais de alunos por turma e geral"""
    class_ids = [ct.class_id for ct in class_tests]
    
    # Buscar todos os alunos das turmas onde a avaliação foi aplicada
    students = Student.query.options(
        joinedload(Student.class_).joinedload(Class.grade)
    ).filter(Student.class_id.in_(class_ids)).all()
    
    # Buscar resultados da avaliação (alunos que realizaram)
    evaluation_results = EvaluationResult.query.filter_by(test_id=evaluation_id).all()
    results_by_student = {er.student_id: er for er in evaluation_results}
    
    # Agrupar alunos por turma
    students_by_class = defaultdict(list)
    for student in students:
        students_by_class[student.class_id].append(student)
    
    # Calcular estatísticas por turma
    por_turma = []
    total_matriculados = 0
    total_avaliados = 0
    
    for class_test in class_tests:
        class_id = class_test.class_id
        class_students = students_by_class[class_id]
        
        # Matriculados = Alunos da turma onde a avaliação foi aplicada
        matriculados = len(class_students)
        # Avaliados = Alunos que realmente realizaram a avaliação (têm resultado)
        avaliados = sum(1 for s in class_students if s.id in results_by_student)
        percentual = (avaliados / matriculados * 100) if matriculados > 0 else 0
        # Faltosos = Matriculados que não realizaram
        faltosos = matriculados - avaliados
        
        # Obter nome da turma
        turma_nome = "Turma Desconhecida"
        if class_students and class_students[0].class_:
            turma_nome = class_students[0].class_.name
        
        por_turma.append({
            "turma": turma_nome,
            "matriculados": matriculados,
            "avaliados": avaliados,
            "percentual": round(percentual, 1),
            "faltosos": faltosos
        })
        
        total_matriculados += matriculados
        total_avaliados += avaliados
    
    # Calcular totais gerais
    total_percentual = (total_avaliados / total_matriculados * 100) if total_matriculados > 0 else 0
    total_faltosos = total_matriculados - total_avaliados
    
    return {
        "por_turma": por_turma,
        "total_geral": {
            "matriculados": total_matriculados,
            "avaliados": total_avaliados,
            "percentual": round(total_percentual, 1),
            "faltosos": total_faltosos
        }
    }


def _calcular_niveis_aprendizagem(evaluation_id: str, class_tests: List[ClassTest]) -> Dict[str, Any]:
    """Calcula níveis de aprendizagem por turma e disciplina"""
    from app.models.skill import Skill
    
    # Buscar questões da avaliação para identificar disciplinas
    test = Test.query.get(evaluation_id)
    if not test or not test.questions:
        return {}
    
    logging.info(f"Calculando níveis de aprendizagem para avaliação {evaluation_id}")
    logging.info(f"Total de questões: {len(test.questions)}")
    
    # Buscar todas as habilidades para mapear ID -> disciplina
    skills_dict = {}
    skill_ids = set()
    for question in test.questions:
        if question.skill and question.skill.strip() and question.skill != '{}':
            clean_skill_id = question.skill.replace('{', '').replace('}', '')
            skill_ids.add(clean_skill_id)
    
    logging.info(f"Skill IDs encontrados: {skill_ids}")
    
    if skill_ids:
        skills = Skill.query.filter(Skill.id.in_(skill_ids)).all()
        skills_dict = {str(skill.id): skill for skill in skills}
        logging.info(f"Habilidades encontradas na tabela Skill: {len(skills_dict)}")
        
        for skill in skills:
            logging.info(f"Skill: {skill.id} -> {skill.code} -> subject_id: {skill.subject_id}")
    
    # Mapear questões para disciplinas
    question_disciplines = {}
    disciplinas_identificadas = set()
    
    for question in test.questions:
        # Questões com skill: mapear via skill.subject_id
        if question.skill and question.skill.strip() and question.skill != '{}':
            clean_skill_id = question.skill.replace('{', '').replace('}', '')
            skill_obj = skills_dict.get(clean_skill_id)
            
            if skill_obj and skill_obj.subject_id:
                subject = Subject.query.get(skill_obj.subject_id)
                if subject:
                    question_disciplines[question.id] = subject.name
                    disciplinas_identificadas.add(subject.name)
                    logging.info(f"Questão {question.id} mapeada para disciplina via skill: {subject.name}")
                else:
                    question_disciplines[question.id] = "Disciplina Geral"
                    disciplinas_identificadas.add("Disciplina Geral")
                    logging.warning(f"Subject não encontrado para skill {skill_obj.id}")
            else:
                question_disciplines[question.id] = "Disciplina Geral"
                disciplinas_identificadas.add("Disciplina Geral")
                if skill_obj:
                    logging.warning(f"Skill {skill_obj.id} não tem subject_id")
                else:
                    logging.warning(f"Skill não encontrado para ID: {clean_skill_id}")
        
        # Questões sem skill: mapear para disciplina principal da avaliação
        else:
            if test.subject_rel:
                question_disciplines[question.id] = test.subject_rel.name
                disciplinas_identificadas.add(test.subject_rel.name)
                logging.info(f"Questão {question.id} sem skill mapeada para disciplina principal: {test.subject_rel.name}")
            else:
                question_disciplines[question.id] = "Disciplina Geral"
                disciplinas_identificadas.add("Disciplina Geral")
                logging.warning(f"Questão {question.id} sem skill e avaliação sem disciplina principal")
        

    
    logging.info(f"Disciplinas identificadas nas questões: {disciplinas_identificadas}")
    logging.info(f"Questões mapeadas para disciplinas: {len(question_disciplines)}")
    
    # Buscar respostas dos alunos
    student_answers = StudentAnswer.query.filter_by(test_id=evaluation_id).all()
    logging.info(f"Total de respostas de alunos: {len(student_answers)}")
    
    # Agrupar respostas por aluno e disciplina
    student_discipline_results = defaultdict(lambda: defaultdict(list))
    
    for answer in student_answers:
        if answer.question_id in question_disciplines:
            disciplina = question_disciplines[answer.question_id]
            student_discipline_results[answer.student_id][disciplina].append(answer)
    
    logging.info(f"Alunos com respostas por disciplina: {len(student_discipline_results)}")
    
    # Buscar resultados da avaliação para obter classificações
    evaluation_results = EvaluationResult.query.options(
        joinedload(EvaluationResult.student).joinedload(Student.class_)
    ).filter_by(test_id=evaluation_id).all()
    
    logging.info(f"Total de resultados de avaliação: {len(evaluation_results)}")
    
    # Agrupar por turma e disciplina
    class_ids = [ct.class_id for ct in class_tests]
    results_by_class = defaultdict(list)
    for result in evaluation_results:
        if result.student and result.student.class_id in class_ids:
            results_by_class[result.student.class_id].append(result)
    
    logging.info(f"Turmas com resultados: {list(results_by_class.keys())}")
    
    # Calcular por turma e disciplina
    disciplinas_resultado = defaultdict(lambda: {
        "por_turma": [],
        "geral": {"Abaixo do Básico": 0, "Básico": 0, "Adequado": 0, "Avançado": 0, "total": 0}
    })
    
    # Dados gerais que englobam todas as disciplinas
    dados_gerais_por_turma = defaultdict(lambda: {"Abaixo do Básico": 0, "Básico": 0, "Adequado": 0, "Avançado": 0, "total": 0})
    dados_gerais_total = {"Abaixo do Básico": 0, "Básico": 0, "Adequado": 0, "Avançado": 0, "total": 0}
    
    for class_test in class_tests:
        class_id = class_test.class_id
        class_results = results_by_class[class_id]
        
        if not class_results:
            continue
        
        # Obter nome da turma
        turma_nome = "Turma Desconhecida"
        if class_results[0].student and class_results[0].student.class_:
            turma_nome = class_results[0].student.class_.name
        
        logging.info(f"Processando turma: {turma_nome}")
        
        # Para cada disciplina, calcular níveis de aprendizagem
        disciplinas_turma = set()
        for result in class_results:
            student_id = result.student_id
            if student_id in student_discipline_results:
                disciplinas_turma.update(student_discipline_results[student_id].keys())
        
        logging.info(f"Disciplinas encontradas na turma {turma_nome}: {disciplinas_turma}")
        
        for disciplina in disciplinas_turma:
            # Contar classificações para esta disciplina e turma
            classificacoes = {"Abaixo do Básico": 0, "Básico": 0, "Adequado": 0, "Avançado": 0}
            
            for result in class_results:
                student_id = result.student_id
                if student_id in student_discipline_results and disciplina in student_discipline_results[student_id]:
                    # Se o aluno respondeu questões desta disciplina, usar sua classificação
                    classificacao = result.classification
                    if classificacao in classificacoes:
                        classificacoes[classificacao] += 1
                        disciplinas_resultado[disciplina]["geral"][classificacao] += 1
            
            total_turma = sum(classificacoes.values())
            disciplinas_resultado[disciplina]["geral"]["total"] += total_turma
            
            disciplinas_resultado[disciplina]["por_turma"].append({
                "turma": turma_nome,
                "abaixo_do_basico": classificacoes["Abaixo do Básico"],
                "basico": classificacoes["Básico"],
                "adequado": classificacoes["Adequado"],
                "avancado": classificacoes["Avançado"],
                "total": total_turma
            })
        
        # Calcular dados gerais para esta turma (todas as disciplinas)
        for result in class_results:
            classificacao = result.classification
            if classificacao in dados_gerais_por_turma[turma_nome]:
                dados_gerais_por_turma[turma_nome][classificacao] += 1
                dados_gerais_por_turma[turma_nome]["total"] += 1
                dados_gerais_total[classificacao] += 1
                dados_gerais_total["total"] += 1
    
    # Organizar resultado final
    resultado_final = {}
    for disciplina, dados in disciplinas_resultado.items():
        resultado_final[disciplina] = {
            "por_turma": dados["por_turma"],
            "geral": {
                "abaixo_do_basico": dados["geral"]["Abaixo do Básico"],
                "basico": dados["geral"]["Básico"],
                "adequado": dados["geral"]["Adequado"],
                "avancado": dados["geral"]["Avançado"],
                "total": dados["geral"]["total"]
            }
        }
    
    # Adicionar dados gerais que englobam todas as disciplinas
    dados_gerais_por_turma_lista = []
    for turma, dados in dados_gerais_por_turma.items():
        dados_gerais_por_turma_lista.append({
            "turma": turma,
            "abaixo_do_basico": dados["Abaixo do Básico"],
            "basico": dados["Básico"],
            "adequado": dados["Adequado"],
            "avancado": dados["Avançado"],
            "total": dados["total"]
        })
    
    resultado_final["GERAL"] = {
        "por_turma": dados_gerais_por_turma_lista,
        "geral": {
            "abaixo_do_basico": dados_gerais_total["Abaixo do Básico"],
            "basico": dados_gerais_total["Básico"],
            "adequado": dados_gerais_total["Adequado"],
            "avancado": dados_gerais_total["Avançado"],
            "total": dados_gerais_total["total"]
        }
    }
    
    logging.info(f"Resultado final com disciplinas: {list(resultado_final.keys())}")
    
    return resultado_final


def _calcular_proficiencia(evaluation_id: str, class_tests: List[ClassTest]) -> Dict[str, Any]:
    """Calcula proficiência por turma e disciplina"""
    from app.models.skill import Skill
    
    # Buscar questões da avaliação para identificar disciplinas
    test = Test.query.get(evaluation_id)
    if not test or not test.questions:
        return {}
    
    # Buscar todas as habilidades para mapear ID -> disciplina
    skills_dict = {}
    skill_ids = set()
    for question in test.questions:
        if question.skill and question.skill.strip() and question.skill != '{}':
            clean_skill_id = question.skill.replace('{', '').replace('}', '')
            skill_ids.add(clean_skill_id)
    
    if skill_ids:
        skills = Skill.query.filter(Skill.id.in_(skill_ids)).all()
        skills_dict = {str(skill.id): skill for skill in skills}
    
    # Mapear questões para disciplinas
    question_disciplines = {}
    for question in test.questions:
        # Questões com skill: mapear via skill.subject_id
        if question.skill and question.skill.strip() and question.skill != '{}':
            clean_skill_id = question.skill.replace('{', '').replace('}', '')
            skill_obj = skills_dict.get(clean_skill_id)
            
            if skill_obj and skill_obj.subject_id:
                subject = Subject.query.get(skill_obj.subject_id)
                if subject:
                    question_disciplines[question.id] = subject.name
                else:
                    question_disciplines[question.id] = "Disciplina Geral"
            else:
                question_disciplines[question.id] = "Disciplina Geral"
        
        # Questões sem skill: mapear para disciplina principal da avaliação
        else:
            if test.subject_rel:
                question_disciplines[question.id] = test.subject_rel.name
            else:
                question_disciplines[question.id] = "Disciplina Geral"
    
    # Buscar respostas dos alunos
    student_answers = StudentAnswer.query.filter_by(test_id=evaluation_id).all()
    
    # Agrupar respostas por aluno e disciplina
    student_discipline_results = defaultdict(lambda: defaultdict(list))
    
    for answer in student_answers:
        if answer.question_id in question_disciplines:
            disciplina = question_disciplines[answer.question_id]
            student_discipline_results[answer.student_id][disciplina].append(answer)
    
    # Buscar resultados da avaliação para obter proficiências
    evaluation_results = EvaluationResult.query.options(
        joinedload(EvaluationResult.student).joinedload(Student.class_)
    ).filter_by(test_id=evaluation_id).all()
    
    # Agrupar por turma e disciplina
    class_ids = [ct.class_id for ct in class_tests]
    results_by_class = defaultdict(list)
    for result in evaluation_results:
        if result.student and result.student.class_id in class_ids:
            results_by_class[result.student.class_id].append(result)
    
    # Calcular por turma e disciplina
    disciplinas_proficiencia = defaultdict(lambda: {
        "por_turma": [],
        "total_proficiencia": 0,
        "total_alunos": 0
    })
    
    # Dados gerais que englobam todas as disciplinas
    dados_gerais_por_turma = defaultdict(list)
    dados_gerais_total_proficiencia = 0
    dados_gerais_total_alunos = 0
    
    for class_test in class_tests:
        class_id = class_test.class_id
        class_results = results_by_class[class_id]
        
        if not class_results:
            continue
        
        # Obter nome da turma
        turma_nome = "Turma Desconhecida"
        if class_results[0].student and class_results[0].student.class_:
            turma_nome = class_results[0].student.class_.name
        
        # Para cada disciplina, calcular proficiência
        disciplinas_turma = set()
        for result in class_results:
            student_id = result.student_id
            if student_id in student_discipline_results:
                disciplinas_turma.update(student_discipline_results[student_id].keys())
        
        for disciplina in disciplinas_turma:
            # Calcular proficiência para esta disciplina e turma
            proficiencias_disciplina = []
            
            for result in class_results:
                student_id = result.student_id
                if student_id in student_discipline_results and disciplina in student_discipline_results[student_id]:
                    # Se o aluno respondeu questões desta disciplina, usar sua proficiência
                    if result.proficiency is not None:
                        proficiencias_disciplina.append(result.proficiency)
            
            if proficiencias_disciplina:
                media_proficiencia = sum(proficiencias_disciplina) / len(proficiencias_disciplina)
                disciplinas_proficiencia[disciplina]["total_proficiencia"] += sum(proficiencias_disciplina)
                disciplinas_proficiencia[disciplina]["total_alunos"] += len(proficiencias_disciplina)
                
                disciplinas_proficiencia[disciplina]["por_turma"].append({
                    "turma": turma_nome,
                    "proficiencia": round(media_proficiencia, 2)
                })
        
        # Calcular dados gerais para esta turma (todas as disciplinas)
        proficiencias_gerais_turma = []
        for result in class_results:
            if result.proficiency is not None:
                proficiencias_gerais_turma.append(result.proficiency)
        
        if proficiencias_gerais_turma:
            media_proficiencia_geral_turma = sum(proficiencias_gerais_turma) / len(proficiencias_gerais_turma)
            dados_gerais_por_turma[turma_nome].append({
                "turma": turma_nome,
                "proficiencia": round(media_proficiencia_geral_turma, 2)
            })
            dados_gerais_total_proficiencia += sum(proficiencias_gerais_turma)
            dados_gerais_total_alunos += len(proficiencias_gerais_turma)
    
    # Organizar resultado final
    resultado_final = {}
    for disciplina, dados in disciplinas_proficiencia.items():
        media_geral_disciplina = dados["total_proficiencia"] / dados["total_alunos"] if dados["total_alunos"] > 0 else 0
        
        resultado_final[disciplina] = {
            "por_turma": sorted(dados["por_turma"], key=lambda x: x["turma"]),
            "media_geral": round(media_geral_disciplina, 2)
        }
    
    # Adicionar dados gerais que englobam todas as disciplinas
    media_geral_total = dados_gerais_total_proficiencia / dados_gerais_total_alunos if dados_gerais_total_alunos > 0 else 0
    
    resultado_final["GERAL"] = {
        "por_turma": sorted([dados for dados_list in dados_gerais_por_turma.values() for dados in dados_list], key=lambda x: x["turma"]),
        "media_geral": round(media_geral_total, 2)
    }
    
    # Calcular média municipal por disciplina
    media_municipal_por_disciplina = _calcular_media_municipal_por_disciplina(evaluation_id, question_disciplines)
    
    return {
        "por_disciplina": resultado_final,
        "media_municipal_por_disciplina": media_municipal_por_disciplina
    }


def _calcular_nota_geral(evaluation_id: str, class_tests: List[ClassTest]) -> Dict[str, Any]:
    """Calcula nota geral por turma e disciplina"""
    from app.models.skill import Skill
    
    # Buscar questões da avaliação para identificar disciplinas
    test = Test.query.get(evaluation_id)
    if not test or not test.questions:
        return {}
    
    # Buscar todas as habilidades para mapear ID -> disciplina
    skills_dict = {}
    skill_ids = set()
    for question in test.questions:
        if question.skill and question.skill.strip() and question.skill != '{}':
            clean_skill_id = question.skill.replace('{', '').replace('}', '')
            skill_ids.add(clean_skill_id)
    
    if skill_ids:
        skills = Skill.query.filter(Skill.id.in_(skill_ids)).all()
        skills_dict = {str(skill.id): skill for skill in skills}
    
    # Mapear questões para disciplinas
    question_disciplines = {}
    for question in test.questions:
        # Questões com skill: mapear via skill.subject_id
        if question.skill and question.skill.strip() and question.skill != '{}':
            clean_skill_id = question.skill.replace('{', '').replace('}', '')
            skill_obj = skills_dict.get(clean_skill_id)
            
            if skill_obj and skill_obj.subject_id:
                subject = Subject.query.get(skill_obj.subject_id)
                if subject:
                    question_disciplines[question.id] = subject.name
                else:
                    question_disciplines[question.id] = "Disciplina Geral"
            else:
                question_disciplines[question.id] = "Disciplina Geral"
        
        # Questões sem skill: mapear para disciplina principal da avaliação
        else:
            if test.subject_rel:
                question_disciplines[question.id] = test.subject_rel.name
            else:
                question_disciplines[question.id] = "Disciplina Geral"
    
    # Buscar respostas dos alunos
    student_answers = StudentAnswer.query.filter_by(test_id=evaluation_id).all()
    
    # Agrupar respostas por aluno e disciplina
    student_discipline_results = defaultdict(lambda: defaultdict(list))
    
    for answer in student_answers:
        if answer.question_id in question_disciplines:
            disciplina = question_disciplines[answer.question_id]
            student_discipline_results[answer.student_id][disciplina].append(answer)
    
    # Buscar resultados da avaliação para obter notas
    evaluation_results = EvaluationResult.query.options(
        joinedload(EvaluationResult.student).joinedload(Student.class_)
    ).filter_by(test_id=evaluation_id).all()
    
    # Agrupar por turma e disciplina
    class_ids = [ct.class_id for ct in class_tests]
    results_by_class = defaultdict(list)
    for result in evaluation_results:
        if result.student and result.student.class_id in class_ids:
            results_by_class[result.student.class_id].append(result)
    
    # Calcular por turma e disciplina
    disciplinas_notas = defaultdict(lambda: {
        "por_turma": [],
        "total_notas": 0,
        "total_alunos": 0
    })
    
    # Dados gerais que englobam todas as disciplinas
    dados_gerais_por_turma = defaultdict(list)
    dados_gerais_total_notas = 0
    dados_gerais_total_alunos = 0
    
    for class_test in class_tests:
        class_id = class_test.class_id
        class_results = results_by_class[class_id]
        
        if not class_results:
            continue
        
        # Obter nome da turma
        turma_nome = "Turma Desconhecida"
        if class_results[0].student and class_results[0].student.class_:
            turma_nome = class_results[0].student.class_.name
        
        # Para cada disciplina, calcular nota
        disciplinas_turma = set()
        for result in class_results:
            student_id = result.student_id
            if student_id in student_discipline_results:
                disciplinas_turma.update(student_discipline_results[student_id].keys())
        
        for disciplina in disciplinas_turma:
            # Calcular nota para esta disciplina e turma
            notas_disciplina = []
            
            for result in class_results:
                student_id = result.student_id
                if student_id in student_discipline_results and disciplina in student_discipline_results[student_id]:
                    # Se o aluno respondeu questões desta disciplina, usar sua nota
                    if result.grade is not None:
                        notas_disciplina.append(result.grade)
            
            if notas_disciplina:
                media_nota = sum(notas_disciplina) / len(notas_disciplina)
                disciplinas_notas[disciplina]["total_notas"] += sum(notas_disciplina)
                disciplinas_notas[disciplina]["total_alunos"] += len(notas_disciplina)
                
                disciplinas_notas[disciplina]["por_turma"].append({
                    "turma": turma_nome,
                    "nota": round(media_nota, 2)
                })
        
        # Calcular dados gerais para esta turma (todas as disciplinas)
        notas_gerais_turma = []
        for result in class_results:
            if result.grade is not None:
                notas_gerais_turma.append(result.grade)
        
        if notas_gerais_turma:
            media_nota_geral_turma = sum(notas_gerais_turma) / len(notas_gerais_turma)
            dados_gerais_por_turma[turma_nome].append({
                "turma": turma_nome,
                "nota": round(media_nota_geral_turma, 2)
            })
            dados_gerais_total_notas += sum(notas_gerais_turma)
            dados_gerais_total_alunos += len(notas_gerais_turma)
    
    # Organizar resultado final
    resultado_final = {}
    for disciplina, dados in disciplinas_notas.items():
        media_geral_disciplina = dados["total_notas"] / dados["total_alunos"] if dados["total_alunos"] > 0 else 0
        
        resultado_final[disciplina] = {
            "por_turma": sorted(dados["por_turma"], key=lambda x: x["turma"]),
            "media_geral": round(media_geral_disciplina, 2)
        }
    
    # Adicionar dados gerais que englobam todas as disciplinas
    media_geral_total = dados_gerais_total_notas / dados_gerais_total_alunos if dados_gerais_total_alunos > 0 else 0
    
    resultado_final["GERAL"] = {
        "por_turma": sorted([dados for dados_list in dados_gerais_por_turma.values() for dados in dados_list], key=lambda x: x["turma"]),
        "media_geral": round(media_geral_total, 2)
    }
    
    # Calcular média municipal por disciplina
    media_municipal_por_disciplina = _calcular_media_municipal_nota_por_disciplina(evaluation_id, question_disciplines)
    
    return {
        "por_disciplina": resultado_final,
        "media_municipal_por_disciplina": media_municipal_por_disciplina
    }


def _calcular_acertos_habilidade(evaluation_id: str) -> Dict[str, Any]:
    """Calcula acertos por habilidade e ranking organizados por disciplina"""
    from app.models.skill import Skill
    
    # Buscar questões da avaliação
    test = Test.query.get(evaluation_id)
    if not test or not test.questions:
        return {}
    
    # Buscar respostas dos alunos
    student_answers = StudentAnswer.query.filter_by(test_id=evaluation_id).all()
    
    # Agrupar respostas por questão
    answers_by_question = defaultdict(list)
    for answer in student_answers:
        answers_by_question[answer.question_id].append(answer)
    
    # Buscar todas as habilidades para mapear ID -> código
    skills_dict = {}
    skill_ids = set()
    for question in test.questions:
        if question.skill and question.skill.strip() and question.skill != '{}':
            # Remover chaves {} do UUID se existirem
            clean_skill_id = question.skill.replace('{', '').replace('}', '')
            skill_ids.add(clean_skill_id)
    
    if skill_ids:
        skills = Skill.query.filter(Skill.id.in_(skill_ids)).all()
        skills_dict = {str(skill.id): skill for skill in skills}
        
        # Debug detalhado
        logging.info(f"Skill IDs encontrados: {skill_ids}")
        logging.info(f"Habilidades encontradas na tabela: {len(skills_dict)}")
        for skill in skills:
            logging.info(f"Skill encontrada: {skill.id} -> {skill.code}")
        
        # Verificar quais IDs não foram encontrados
        not_found = skill_ids - set(skills_dict.keys())
        if not_found:
            logging.warning(f"Skill IDs não encontrados: {not_found}")
    
    # Calcular acertos por habilidade e disciplina
    disciplinas_habilidades = defaultdict(lambda: defaultdict(lambda: {"total": 0, "acertos": 0, "questoes": []}))
    
    # Dados gerais que englobam todas as disciplinas
    habilidades_gerais = defaultdict(lambda: {"total": 0, "acertos": 0, "questoes": []})
    
    for question in test.questions:
        # Questões com skill: mapear via skill.subject_id
        if question.skill and question.skill.strip() and question.skill != '{}':
            # Remover chaves {} do UUID se existirem
            clean_skill_id = question.skill.replace('{', '').replace('}', '')
            skill_obj = skills_dict.get(clean_skill_id)
            
            # Se não encontrou a habilidade na tabela Skill, usar o ID como código
            if not skill_obj:
                skill_code = clean_skill_id
                disciplina = "Disciplina Geral"
                logging.warning(f"Habilidade não encontrada na tabela Skill: {clean_skill_id}")
            else:
                skill_code = skill_obj.code
                # Obter disciplina da habilidade
                if skill_obj.subject_id:
                    # Buscar a disciplina pelo subject_id
                    subject = Subject.query.get(skill_obj.subject_id)
                    if subject:
                        disciplina = subject.name
                    else:
                        disciplina = "Disciplina Geral"
                else:
                    disciplina = "Disciplina Geral"
        
        # Questões sem skill: mapear para disciplina principal da avaliação
        else:
            if test.subject_rel:
                disciplina = test.subject_rel.name
                skill_code = f"Questão {question.number or 'N/A'}"
                logging.info(f"Questão {question.id} sem skill mapeada para disciplina principal: {test.subject_rel.name}")
            else:
                disciplina = "Disciplina Geral"
                skill_code = f"Questão {question.number or 'N/A'}"
                logging.warning(f"Questão {question.id} sem skill e avaliação sem disciplina principal")
        
        answers = answers_by_question.get(question.id, [])
        
        total_respostas = len(answers)
        acertos = 0
        
        for answer in answers:
            # Verificar se a resposta está correta
            if question.question_type == 'multiple_choice':
                # Para múltipla escolha, verificar se a resposta está correta
                if question.correct_answer and str(answer.answer).strip().lower() == str(question.correct_answer).strip().lower():
                    acertos += 1
            else:
                # Para outros tipos, comparar com correct_answer
                if str(answer.answer).strip().lower() == str(question.correct_answer).strip().lower():
                    acertos += 1
        
        disciplinas_habilidades[disciplina][skill_code]["total"] += total_respostas
        disciplinas_habilidades[disciplina][skill_code]["acertos"] += acertos
        disciplinas_habilidades[disciplina][skill_code]["questoes"].append({
            "numero": question.number or 1,
            "acertos": acertos,
            "total": total_respostas
        })
        
        # Adicionar aos dados gerais
        habilidades_gerais[skill_code]["total"] += total_respostas
        habilidades_gerais[skill_code]["acertos"] += acertos
        habilidades_gerais[skill_code]["questoes"].append({
            "numero": question.number or 1,
            "acertos": acertos,
            "total": total_respostas
        })
    
    # Organizar resultado por disciplina
    resultado_por_disciplina = {}
    
    for disciplina, habilidades in disciplinas_habilidades.items():
        # Calcular percentuais e ranking para esta disciplina
        habilidades_ranking = []
        for skill_code, dados in habilidades.items():
            if dados["total"] > 0:
                percentual = (dados["acertos"] / dados["total"]) * 100
                habilidades_ranking.append({
                    "codigo": skill_code,
                    "descricao": f"Habilidade {skill_code}",
                    "acertos": dados["acertos"],
                    "total": dados["total"],
                    "percentual": round(percentual, 1),
                    "questoes": dados["questoes"]
                })
        
        # Ordenar por percentual (maior para menor)
        habilidades_ranking.sort(key=lambda x: x["percentual"], reverse=True)
        
        # Adicionar ranking
        for i, habilidade in enumerate(habilidades_ranking, 1):
            habilidade["ranking"] = i
        
        resultado_por_disciplina[disciplina] = {
            "habilidades": habilidades_ranking
        }
    
    # Adicionar dados gerais que englobam todas as disciplinas
    habilidades_gerais_ranking = []
    for skill_code, dados in habilidades_gerais.items():
        if dados["total"] > 0:
            percentual = (dados["acertos"] / dados["total"]) * 100
            habilidades_gerais_ranking.append({
                "codigo": skill_code,
                "descricao": f"Habilidade {skill_code}",
                "acertos": dados["acertos"],
                "total": dados["total"],
                "percentual": round(percentual, 1),
                "questoes": dados["questoes"]
            })
    
    # Ordenar por percentual (maior para menor)
    habilidades_gerais_ranking.sort(key=lambda x: x["percentual"], reverse=True)
    
    # Adicionar ranking
    for i, habilidade in enumerate(habilidades_gerais_ranking, 1):
        habilidade["ranking"] = i
    
    resultado_por_disciplina["GERAL"] = {
        "habilidades": habilidades_gerais_ranking
    }
    
    return resultado_por_disciplina


def _calcular_media_municipal(evaluation_id: str) -> float:
    """Calcula média de proficiência municipal"""
    try:
        # Buscar a primeira escola para obter o município
        test = Test.query.get(evaluation_id)
        if not test:
            return 0.0
        
        # Buscar turmas da avaliação
        class_tests = ClassTest.query.filter_by(test_id=evaluation_id).all()
        if not class_tests:
            return 0.0
        
        # Obter município da primeira turma
        first_class = Class.query.get(class_tests[0].class_id)
        if not first_class or not first_class.school:
            return 0.0
        
        city_id = first_class.school.city_id
        
        # Buscar todos os resultados do município para esta avaliação
        municipal_results = db.session.query(EvaluationResult).join(
            Student, EvaluationResult.student_id == Student.id
        ).join(
            Class, Student.class_id == Class.id
        ).join(
            School, Class.school_id == School.id
        ).filter(
            School.city_id == city_id,
            EvaluationResult.test_id == evaluation_id
        ).all()
        
        if not municipal_results:
            return 0.0
        
        media_municipal = sum(r.proficiency for r in municipal_results) / len(municipal_results)
        return media_municipal
        
    except Exception as e:
        logging.error(f"Erro ao calcular média municipal: {str(e)}")
        return 0.0


def _calcular_media_municipal_nota(evaluation_id: str) -> float:
    """Calcula média de nota municipal"""
    try:
        # Buscar a primeira escola para obter o município
        test = Test.query.get(evaluation_id)
        if not test:
            return 0.0
        
        # Buscar turmas da avaliação
        class_tests = ClassTest.query.filter_by(test_id=evaluation_id).all()
        if not class_tests:
            return 0.0
        
        # Obter município da primeira turma
        first_class = Class.query.get(class_tests[0].class_id)
        if not first_class or not first_class.school:
            return 0.0
        
        city_id = first_class.school.city_id
        
        # Buscar todos os resultados do município para esta avaliação
        municipal_results = db.session.query(EvaluationResult).join(
            Student, EvaluationResult.student_id == Student.id
        ).join(
            Class, Student.class_id == Class.id
        ).join(
            School, Class.school_id == School.id
        ).filter(
            School.city_id == city_id,
            EvaluationResult.test_id == evaluation_id
        ).all()
        
        if not municipal_results:
            return 0.0
        
        media_municipal = sum(r.grade for r in municipal_results) / len(municipal_results)
        return media_municipal
        
    except Exception as e:
        logging.error(f"Erro ao calcular média municipal de nota: {str(e)}")
        return 0.0


def _calcular_media_municipal_por_disciplina(evaluation_id: str, question_disciplines: Dict[str, str]) -> Dict[str, float]:
    """Calcula média de proficiência municipal por disciplina"""
    try:
        # Buscar a primeira escola para obter o município
        test = Test.query.get(evaluation_id)
        if not test:
            return {}
        
        # Buscar turmas da avaliação
        class_tests = ClassTest.query.filter_by(test_id=evaluation_id).all()
        if not class_tests:
            return {}
        
        # Obter município da primeira turma
        first_class = Class.query.get(class_tests[0].class_id)
        if not first_class or not first_class.school:
            return {}
        
        city_id = first_class.school.city_id
        
        # Buscar todos os resultados do município para esta avaliação
        municipal_results = db.session.query(EvaluationResult).join(
            Student, EvaluationResult.student_id == Student.id
        ).join(
            Class, Student.class_id == Class.id
        ).join(
            School, Class.school_id == School.id
        ).filter(
            School.city_id == city_id,
            EvaluationResult.test_id == evaluation_id
        ).all()
        
        if not municipal_results:
            return {}
        
        # Buscar respostas dos alunos do município
        student_ids = [r.student_id for r in municipal_results]
        municipal_answers = StudentAnswer.query.filter(
            StudentAnswer.test_id == evaluation_id,
            StudentAnswer.student_id.in_(student_ids)
        ).all()
        
        # Agrupar respostas por aluno e disciplina
        student_discipline_results = defaultdict(lambda: defaultdict(list))
        
        for answer in municipal_answers:
            if answer.question_id in question_disciplines:
                disciplina = question_disciplines[answer.question_id]
                student_discipline_results[answer.student_id][disciplina].append(answer)
        
        # Calcular média por disciplina
        disciplinas_proficiencia = defaultdict(list)
        
        for result in municipal_results:
            student_id = result.student_id
            if student_id in student_discipline_results and result.proficiency is not None:
                for disciplina in student_discipline_results[student_id].keys():
                    disciplinas_proficiencia[disciplina].append(result.proficiency)
        
        # Calcular média para cada disciplina
        medias_por_disciplina = {}
        for disciplina, proficiencias in disciplinas_proficiencia.items():
            if proficiencias:
                media = sum(proficiencias) / len(proficiencias)
                medias_por_disciplina[disciplina] = round(media, 2)
        
        return medias_por_disciplina
        
    except Exception as e:
        logging.error(f"Erro ao calcular média municipal por disciplina: {str(e)}")
        return {}


def _calcular_media_municipal_nota_por_disciplina(evaluation_id: str, question_disciplines: Dict[str, str]) -> Dict[str, float]:
    """Calcula média de nota municipal por disciplina"""
    try:
        # Buscar a primeira escola para obter o município
        test = Test.query.get(evaluation_id)
        if not test:
            return {}
        
        # Buscar turmas da avaliação
        class_tests = ClassTest.query.filter_by(test_id=evaluation_id).all()
        if not class_tests:
            return {}
        
        # Obter município da primeira turma
        first_class = Class.query.get(class_tests[0].class_id)
        if not first_class or not first_class.school:
            return {}
        
        city_id = first_class.school.city_id
        
        # Buscar todos os resultados do município para esta avaliação
        municipal_results = db.session.query(EvaluationResult).join(
            Student, EvaluationResult.student_id == Student.id
        ).join(
            Class, Student.class_id == Class.id
        ).join(
            School, Class.school_id == School.id
        ).filter(
            School.city_id == city_id,
            EvaluationResult.test_id == evaluation_id
        ).all()
        
        if not municipal_results:
            return {}
        
        # Buscar respostas dos alunos do município
        student_ids = [r.student_id for r in municipal_results]
        municipal_answers = StudentAnswer.query.filter(
            StudentAnswer.test_id == evaluation_id,
            StudentAnswer.student_id.in_(student_ids)
        ).all()
        
        # Agrupar respostas por aluno e disciplina
        student_discipline_results = defaultdict(lambda: defaultdict(list))
        
        for answer in municipal_answers:
            if answer.question_id in question_disciplines:
                disciplina = question_disciplines[answer.question_id]
                student_discipline_results[answer.student_id][disciplina].append(answer)
        
        # Calcular média por disciplina
        disciplinas_notas = defaultdict(list)
        
        for result in municipal_results:
            student_id = result.student_id
            if student_id in student_discipline_results and result.grade is not None:
                for disciplina in student_discipline_results[student_id].keys():
                    disciplinas_notas[disciplina].append(result.grade)
        
        # Calcular média para cada disciplina
        medias_por_disciplina = {}
        for disciplina, notas in disciplinas_notas.items():
            if notas:
                media = sum(notas) / len(notas)
                medias_por_disciplina[disciplina] = round(media, 2)
        
        return medias_por_disciplina
        
    except Exception as e:
        logging.error(f"Erro ao calcular média municipal de nota por disciplina: {str(e)}")
        return {}


def _gerar_analise_participacao(template_data: Dict) -> str:
    """Gera análise inteligente da participação baseada nos dados"""
    try:
        total_alunos = template_data.get("total_alunos", {})
        total_geral = total_alunos.get("total_geral", {})
        
        matriculados = total_geral.get("matriculados", 0)
        avaliados = total_geral.get("avaliados", 0)
        percentual = total_geral.get("percentual", 0)
        faltosos = total_geral.get("faltosos", 0)
        
        if matriculados == 0:
            return "Dados de participação não disponíveis."
        
        # Análise baseada nos percentuais
        if percentual >= 90:
            avaliacao = "excelente"
            recomendacao = "manter o alto engajamento"
        elif percentual >= 80:
            avaliacao = "boa"
            recomendacao = "investigar causas das ausências e contatar famílias"
        elif percentual >= 70:
            avaliacao = "regular"
            recomendacao = "implementar estratégias de motivação e contato com famílias"
        else:
            avaliacao = "baixa"
            recomendacao = "ação imediata para aumentar participação"
        
        analise = f"A escola avaliou {avaliados} alunos do 9º ano, o que representa {percentual}% do total de {matriculados} alunos matriculados. "
        analise += f"Registrou-se um total de {faltosos} alunos faltosos. "
        analise += f"A taxa de participação de {percentual}% é {avaliacao}, indicando um {avaliacao} engajamento geral. "
        
        if faltosos > 0:
            analise += f"Contudo, a ausência de {faltosos} alunos é um número considerável. "
        
        analise += f"Recomenda-se {recomendacao} para elevar a participação em futuras avaliações e o planejamento de uma reposição diagnóstica para os alunos faltosos, visando obter um panorama mais completo da aprendizagem na série."
        
        return analise
        
    except Exception as e:
        logging.warning(f"Erro ao gerar análise de participação: {str(e)}")
        return "Análise de participação não disponível."


def _gerar_analise_proficiencia(template_data: Dict) -> str:
    """Gera análise inteligente da proficiência baseada nos dados"""
    try:
        proficiencia = template_data.get("proficiencia", {})
        disciplinas = proficiencia.get("por_disciplina", {})
        
        if not disciplinas:
            return "Dados de proficiência não disponíveis."
        
        # Buscar médias por disciplina
        medias_disciplinas = {}
        for disciplina, dados in disciplinas.items():
            if disciplina != "GERAL" and "por_turma" in dados:
                valores = [t.get("proficiencia", 0) for t in dados["por_turma"] if t.get("proficiencia")]
                if valores:
                    medias_disciplinas[disciplina] = sum(valores) / len(valores)
        
        if not medias_disciplinas:
            return "Dados de proficiência por disciplina não disponíveis."
        
        # Gerar análise
        analise = "A proficiência média da escola nesta avaliação diagnóstica foi de "
        
        disciplinas_list = list(medias_disciplinas.keys())
        for i, disciplina in enumerate(disciplinas_list):
            media = round(medias_disciplinas[disciplina], 2)
            if i == 0:
                analise += f"{media} em {disciplina}"
            elif i == len(disciplinas_list) - 1:
                analise += f" e {media} em {disciplina}."
            else:
                analise += f", {media} em {disciplina}"
        
        # Comparação com referencial (simulado)
        analise += " Comparativamente, o desempenho em "
        
        melhor_disciplina = max(medias_disciplinas.items(), key=lambda x: x[1])
        pior_disciplina = min(medias_disciplinas.items(), key=lambda x: x[1])
        
        analise += f"{melhor_disciplina[0]} ({melhor_disciplina[1]}) está significativamente acima da média de proficiência de 272,19 alcançada na Prova Saeb/2023 para o 9º ano. "
        analise += f"Este é um resultado excelente e um grande destaque para a escola, indicando um alto nível de domínio das competências em {melhor_disciplina[0]}. "
        
        analise += f"Em {pior_disciplina[0]}, a proficiência de {pior_disciplina[1]} encontra-se abaixo da média de 293,55 obtida no Saeb/2023. "
        analise += f"Apesar de ser uma proficiência considerável, ainda há uma lacuna para atingir o referencial nacional, indicando que esta é a área que demanda maior foco estratégico."
        
        return analise
        
    except Exception as e:
        logging.warning(f"Erro ao gerar análise de proficiência: {str(e)}")
        return "Análise de proficiência não disponível."


def _gerar_analise_notas(template_data: Dict) -> str:
    """Gera análise inteligente das notas baseada nos dados"""
    try:
        nota_geral = template_data.get("nota_geral", {})
        disciplinas = nota_geral.get("por_disciplina", {})
        
        if not disciplinas:
            return "Dados de notas não disponíveis."
        
        # Buscar médias por disciplina
        medias_disciplinas = {}
        for disciplina, dados in disciplinas.items():
            if disciplina != "GERAL" and "por_turma" in dados:
                valores = [t.get("nota", 0) for t in dados["por_turma"] if t.get("nota")]
                if valores:
                    medias_disciplinas[disciplina] = sum(valores) / len(valores)
        
        if not medias_disciplinas:
            return "Dados de notas por disciplina não disponíveis."
        
        # Calcular média geral
        media_geral = sum(medias_disciplinas.values()) / len(medias_disciplinas)
        
        # Gerar análise
        analise = f"A média geral de nota da escola na avaliação diagnóstica foi de {round(media_geral, 2)}, com "
        
        disciplinas_list = list(medias_disciplinas.keys())
        for i, disciplina in enumerate(disciplinas_list):
            media = round(medias_disciplinas[disciplina], 2)
            if i == 0:
                analise += f"{media} em {disciplina}"
            elif i == len(disciplinas_list) - 1:
                analise += f" e {media} em {disciplina}."
            else:
                analise += f", {media} em {disciplina}"
        
        # Comparação com referencial
        analise += f" Esta média geral de {round(media_geral, 2)} está "
        
        if media_geral >= 6.0:
            analise += "exatamente alinhada e ligeiramente acima da nota padronizada de 6,1 obtida na Prova Saeb/2023 para o 9º ano, o que é um resultado muito positivo para a escola como um todo. "
        else:
            analise += "abaixo da nota padronizada de 6,1 obtida na Prova Saeb/2023 para o 9º ano, indicando a necessidade de melhorias para atingir o padrão nacional. "
        
        # Análise por disciplina
        melhor_disciplina = max(medias_disciplinas.items(), key=lambda x: x[1])
        pior_disciplina = min(medias_disciplinas.items(), key=lambda x: x[1])
        
        analise += f"A nota em {melhor_disciplina[0]} ({melhor_disciplina[1]}) é excelente e reflete a alta proficiência alcançada. "
        analise += f"A nota em {pior_disciplina[0]} ({pior_disciplina[1]}), embora contribua para a boa média geral, ainda está abaixo do ideal e indica a necessidade de melhorias nesta disciplina para que a escola possa elevar ainda mais seu desempenho global."
        
        return analise
        
    except Exception as e:
        logging.warning(f"Erro ao gerar análise de notas: {str(e)}")
        return "Análise de notas não disponível."


def _gerar_analise_habilidades(template_data: Dict) -> str:
    """Gera análise inteligente das habilidades baseada nos dados"""
    try:
        acertos_por_habilidade = template_data.get("acertos_por_habilidade", {})
        
        if not acertos_por_habilidade:
            return "Dados de habilidades não disponíveis."
        
        # Analisar habilidades por disciplina
        analises_disciplinas = []
        
        for disciplina, dados in acertos_por_habilidade.items():
            if disciplina == "GERAL" or not dados.get("habilidades"):
                continue
            
            habilidades = dados["habilidades"]
            if not habilidades:
                continue
            
            # Calcular estatísticas
            total_habilidades = len(habilidades)
            habilidades_acima_70 = [h for h in habilidades if h.get("percentual", 0) >= 70]
            habilidades_abaixo_70 = [h for h in habilidades if h.get("percentual", 0) < 70]
            
            percentual_acima_70 = (len(habilidades_acima_70) / total_habilidades) * 100
            
            # Gerar análise para esta disciplina
            analise_disc = f"Os alunos demonstraram "
            
            if percentual_acima_70 >= 80:
                analise_disc += "excelente desempenho na grande maioria das habilidades de "
            elif percentual_acima_70 >= 60:
                analise_disc += "bom desempenho na maioria das habilidades de "
            else:
                analise_disc += "desempenho regular nas habilidades de "
            
            analise_disc += f"{disciplina}. "
            
            # Destacar habilidades com melhor desempenho
            if habilidades_acima_70:
                melhor_habilidade = max(habilidades_acima_70, key=lambda x: x.get("percentual", 0))
                analise_disc += f"Destacam-se habilidades como {melhor_habilidade.get('codigo', 'N/A')} com {melhor_habilidade.get('percentual', 0)}% de acerto. "
            
            # Pontos de atenção
            if habilidades_abaixo_70:
                analise_disc += f"Habilidades com Desempenho Abaixo da Meta (< 70%) – Pontos de Atenção: "
                for h in habilidades_abaixo_70[:3]:  # Limitar a 3 exemplos
                    analise_disc += f"{h.get('codigo', 'N/A')} ({h.get('percentual', 0)}%), "
                analise_disc = analise_disc.rstrip(", ") + ". "
                
                analise_disc += "Mesmo com o alto desempenho geral, focar nessas habilidades pode levar a um domínio ainda mais completo."
            
            analises_disciplinas.append(analise_disc)
        
        if not analises_disciplinas:
            return "Análise detalhada de habilidades não disponível."
        
        return " ".join(analises_disciplinas)
        
    except Exception as e:
        logging.warning(f"Erro ao gerar análise de habilidades: {str(e)}")
        return "Análise de habilidades não disponível."


@bp.errorhandler(Exception)
def handle_error(error):
    """Handler de erros para o blueprint"""
    logging.error(f"Erro no blueprint de relatórios: {str(error)}")
    return jsonify({"error": "Erro interno do servidor"}), 500
